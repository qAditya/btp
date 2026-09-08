from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH as A
doc=Document('_part4.docx')
for s in doc.sections: s.top_margin=Inches(1);s.bottom_margin=Inches(1);s.left_margin=Inches(1.5);s.right_margin=Inches(1)
def H(t,sz=14,al=A.LEFT):
    p=doc.add_paragraph();p.alignment=al;p.paragraph_format.space_before=Pt(12);p.paragraph_format.space_after=Pt(6);r=p.add_run(t);r.bold=True;r.font.size=Pt(sz);r.font.name='Times New Roman'
def P(t,al=A.JUSTIFY):
    p=doc.add_paragraph();p.alignment=al;p.paragraph_format.space_after=Pt(6);p.paragraph_format.space_before=Pt(3);r=p.add_run(t);r.font.size=Pt(12);r.font.name='Times New Roman'
def NL():
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(0);p.paragraph_format.space_after=Pt(0)

H("CHAPTER 5",16,A.CENTER);H("CONCLUSION AND FUTURE WORK",14,A.CENTER);NL()

H("5.1  Summary of Contributions")
P("This project has presented a comprehensive computational framework for the parametric analysis and optimization of bifacial photovoltaic systems. The work addresses the identified gaps in existing literature by providing a systematic, accessible, and location-specific tool for determining optimal installation configurations. The key contributions of this work are summarized below.")
P("First, a rigorous mathematical model for computing the front-side and rear-side irradiance on bifacial PV modules has been implemented. The model incorporates the Liu-Jordan isotropic sky model [15] for front-side irradiance estimation and the Yusufoglu view factor approach [18] for accurate rear-side irradiance calculation with shadow correction. The mathematical formulation accounts for the effects of tilt angle, ground albedo, mounting height, and solar geometry on both front-side and rear-side irradiance contributions, providing a comprehensive representation of bifacial PV module operation.")
P("Second, a MATLAB/Simulink model has been developed for simulating the electrical characteristics (I-V and P-V curves) of bifacial PV modules under varying irradiance and temperature conditions. The Simulink model utilizes a single-diode equivalent circuit representation of the PV module and provides validated electrical output predictions for any combination of input conditions. The model enables visualization of the impact of different installation configurations on the electrical performance of the module.")
P("Third, a user-oriented frontend platform has been designed and implemented that integrates location-based solar irradiance data retrieval from the NASA POWER API with automated parametric sweep analyses. The platform provides an accessible tool for PV system designers, requiring no specialized software knowledge or access to commercial simulation tools. The modular architecture of the platform ensures extensibility and ease of maintenance.")
P("Fourth, the framework supports dual optimization objectives (maximum energy and maximum rear gain), enabling system designers to evaluate trade-offs and make informed configuration decisions. This dual optimization capability addresses a gap in existing tools that typically consider only a single performance metric.")
P("Fifth, the proposed approach has been validated through a case study for Greater Noida, India, with results showing good agreement with published literature. The validation confirms the reliability and accuracy of the mathematical formulation and computational implementation.")

H("5.2  Key Findings")
P("The parametric analyses conducted in this study have yielded the following key findings that provide practical guidance for the design and optimization of bifacial PV installations:")
P("(i) For the case study location (Greater Noida, India, 28.47\u00b0N) with a concrete ground surface (albedo = 0.3), the optimal configuration for maximum energy output was identified as a tilt angle of 30 degrees and mounting height of 450 cm, achieving a total energy of 3.168 kWh and rear gain of 16.55%. The optimal tilt angle closely corresponds to the site latitude, consistent with established PV design principles.")
P("(ii) The optimal configuration for maximum rear gain was a tilt angle of 10 degrees and height of 450 cm, achieving a rear gain of 17.90% but with a lower total energy of 2.994 kWh (approximately 5.5% lower than the maximum energy configuration). This demonstrates the existence of a clear trade-off between total energy maximization and rear gain maximization.")
P("(iii) Ground surface albedo was confirmed as the most influential parameter affecting rear-side irradiance contribution. The rear gain varied from below 5% for low-albedo surfaces (such as asphalt) to over 24% for highly reflective surfaces (such as aluminum), representing a nearly five-fold variation. The relationship between albedo and rear gain was found to be approximately linear, consistent with the mathematical formulation and published literature.")
P("(iv) Mounting height has a positive but diminishing effect on rear-side performance, exhibiting a near-logarithmic dependence consistent with the findings of Yusufoglu et al. [18]. The practical implication is that mounting heights in the range of 200-300 cm provide a good balance between performance improvement and structural feasibility for most applications.")
P("(v) The interaction between tilt angle, height, and albedo necessitates combined parametric optimization rather than independent single-parameter studies. The combined sweep approach captured interaction effects that would be missed by analyzing parameters in isolation, leading to more accurate identification of globally optimal configurations.")
P("(vi) The system performance was found to be relatively robust to small variations near the optimal configuration, with the top five configurations differing by less than 0.25% in total energy. This robustness is practically significant as it allows flexibility in accommodating site-specific constraints without significant performance penalties.")

H("5.3  Limitations")
P("The present work has certain limitations that should be acknowledged and considered when interpreting the results:")
P("(i) The mathematical model assumes uniform ground albedo beneath and surrounding the module, which may not accurately represent real-world installations where ground surface properties can vary spatially due to vegetation, shadows from nearby structures, or mixed surface types.")
P("(ii) The analysis is based on single-day irradiance data. While this approach provides useful insights into the relative performance of different configurations, extending the analysis to annual energy yield computations would provide more comprehensive optimization results that account for seasonal variations in solar geometry and irradiance conditions.")
P("(iii) The current implementation does not account for inter-row shading effects in multi-row array configurations, which can be significant in large-scale installations where adjacent rows cast shadows on the ground and on each other.")
P("(iv) Temperature effects on module performance are modeled using the simplified NOCT approach. More detailed thermal models that account for wind speed, mounting configuration, and rear-surface ventilation may improve the accuracy of energy yield predictions.")
P("(v) The isotropic sky model used for diffuse radiation estimation may underestimate the circumsolar and horizon brightening components that contribute to tilted surface irradiance, particularly under partly cloudy conditions. More advanced anisotropic models such as the Perez model could improve the accuracy of front-side irradiance estimation.")

H("5.4  Scope for Future Work")
P("Several promising directions for future research and development are identified based on the outcomes and limitations of this project:")
P("(i) Extension to Annual Energy Yield: The current framework can be extended to perform year-round energy yield simulations by aggregating daily analyses across all days of the year. This extension would provide more comprehensive optimization guidance that accounts for seasonal variations in solar geometry, temperature, and irradiance conditions, enabling more robust design decisions for long-term system performance.")
P("(ii) Multi-Row Array Modeling: Incorporating inter-row shading models would enable the framework to optimize row-to-row spacing in addition to tilt, height, and albedo. This enhancement is essential for utility-scale bifacial solar farm design where inter-row interactions significantly affect system performance.")
P("(iii) Machine Learning Integration: Developing machine learning models trained on the parametric sweep data could enable rapid prediction of optimal configurations without requiring full computational analysis. Neural network or random forest models could be trained on a comprehensive database of simulation results to provide near-instantaneous optimization guidance for new locations.")
P("(iv) Economic Analysis: Integrating techno-economic analysis capabilities would allow the framework to evaluate the cost-effectiveness of different bifacial PV configurations. This analysis would account for structural costs associated with different mounting heights, ground preparation costs for albedo enhancement, and the economic value of additional energy generation through levelized cost of energy (LCOE) calculations.")
P("(v) Real-Time Monitoring: Developing a real-time monitoring system that continuously compares actual system performance with predicted values could enable adaptive optimization and early detection of performance degradation. Integration with IoT sensors and cloud-based analytics platforms could provide automated performance assessment and maintenance alerts.")
P("(vi) Geographic Information System Integration: Combining the optimization framework with GIS tools could enable automated site assessment for large-scale bifacial PV deployment planning. This integration would incorporate terrain data, land use classification, and local albedo mapping from satellite imagery to provide comprehensive site suitability analysis.")
P("(vii) Tracking System Integration: Extending the framework to support single-axis and dual-axis tracking configurations would enable comparison of fixed-tilt and tracking bifacial PV systems, providing a more complete optimization tool for diverse installation scenarios.")

H("5.5  Task, Achievements and Possible Beneficiaries")
P("The primary task of this project was to develop a computational framework for parametric analysis and optimization of bifacial PV installations. This task has been successfully accomplished through the design and implementation of an integrated platform comprising mathematical modeling, MATLAB/Simulink simulation, and a web-based frontend interface.")
P("The key achievements include: development of a validated mathematical model for bifacial PV irradiance computation; creation of a Simulink model for electrical performance simulation; implementation of a user-friendly frontend platform with NASA POWER API integration; and demonstration of the framework through a comprehensive case study with comparison to published literature.")
P("The possible beneficiaries of this work include: PV system designers and engineers seeking optimal configuration guidance for bifacial installations; solar project developers evaluating the feasibility and expected performance of bifacial PV projects; academic researchers studying bifacial PV technology and requiring a flexible simulation platform; and policymakers and planners assessing the potential of bifacial PV technology for renewable energy targets.")

doc.add_page_break()

# ===================== REFERENCES =====================
H("REFERENCES",16,A.CENTER);NL()
refs=[
'[1] R. Guerrero-Lemus, R. Vega, T. Kim, A. Kimm and L. E. Shephard, "Bifacial solar photovoltaics \u2013 A technology review," Renewable and Sustainable Energy Reviews, vol. 60, pp. 1533-1549, 2016.',
'[2] M. Alam, M. R. Khan and C. Deline, "Performance comparison between bifacial and monofacial photovoltaic systems under varying environmental conditions," Sustainable Energy Technologies and Assessments, vol. 57, p. 103210, 2023.',
'[3] C. Deline, S. A. Pelaez, S. MacAlpine and B. Marion, "Bifacial photovoltaic system performance: Separating fact from fiction," National Renewable Energy Laboratory (NREL), Golden, CO, USA, Tech. Rep. NREL/CP-5K00-74090, 2019.',
'[4] X. Sun, M. R. Khan, C. Deline and M. A. Alam, "Optimization and performance of bifacial solar modules: A global perspective," Applied Energy, vol. 212, pp. 1601-1610, 2018.',
'[5] A. Basak, S. Das and R. Banerjee, "Tilt angle optimization for bifacial photovoltaic modules for enhanced energy generation," Applied Energy, vol. 345, p. 121245, 2025.',
'[6] U. Peter and M. Novak, "Optimal tilt angle for maximizing energy production of bifacial solar panels," Journal of Sustainable Energy Systems, vol. 14, no. 1, pp. 45-56, 2025.',
'[7] N. Riedel-Lyngskjaer, J. M. L. P. Larsen and B. G. Nielsen, "Effect of spectral albedo on energy yield of bifacial photovoltaic systems," Solar Energy, vol. 231, pp. 176-186, 2022.',
'[8] M. H. Aksoy and H. A. Ceylan, "Investigation of ground albedo effects on bifacial photovoltaic panel performance," Renewable Energy Research Journal, vol. 13, no. 4, pp. 1800-1808, 2023.',
'[9] N. Baghel, A. Chandel and S. K. Sharma, "Performance evaluation and optimization of albedo and tilt angle in bifacial photovoltaic systems," Solar Energy, vol. 250, pp. 84-95, 2023.',
'[10] I. N. Atalay, M. Kaya and H. Ozturk, "Experimental investigation of ground surface reflectivity impact on bifacial photovoltaic systems," Solar Energy, vol. 246, pp. 65-74, 2023.',
'[11] F. Dincer and E. Ozer, "Optimization of rear-side energy contribution in bifacial PV panels: A parametric analysis on albedo, tilt, height, and mounting configuration," Energies, vol. 18, no. 17, p. 4443, 2025.',
'[12] D. S. Braga, R. Zilles and E. Lorenzo, "Performance analysis of bifacial photovoltaic modules under tropical climatic conditions," Renewable Energy and Environmental Sustainability, vol. 8, pp. 1-10, 2023.',
'[13] A. F. Almarshoud, A. Al-Badi and M. Al-Sulaiman, "Experimental performance analysis of bifacial photovoltaic modules under different operating conditions," Energies, vol. 17, no. 21, p. 5456, 2024.',
'[14] M. T. Patel, M. R. Khan and M. A. Alam, "Optimum design of tracking bifacial solar farms," IEEE Journal of Photovoltaics, vol. 11, no. 1, pp. 237-245, 2021.',
'[15] B. Y. H. Liu and R. C. Jordan, "The interrelationship and characteristic distribution of direct, diffuse and total solar radiation," Solar Energy, vol. 4, no. 3, pp. 1-19, 1960.',
'[16] J. A. Duffie and W. A. Beckman, Solar Engineering of Thermal Processes, 4th ed. Hoboken, NJ, USA: Wiley, 2013.',
'[17] B. Marion, S. MacAlpine, C. Deline, A. Asgharzadeh, F. Toor, D. Riley, J. Stein and C. Hansen, "A practical irradiance model for bifacial PV modules," in Proc. 44th IEEE Photovoltaic Specialists Conf. (PVSC), Washington, DC, USA, 2017, pp. 1537-1542.',
'[18] U. A. Yusufoglu, T. M. Pletzer, L. J. Koduvelikulathu, C. Comparotto, R. Kopecek and H. Kurz, "Analysis of the annual performance of bifacial modules and optimization methods," IEEE Journal of Photovoltaics, vol. 5, no. 1, pp. 320-328, 2015.',
'[19] P. K. Sahu, J. N. Roy and C. Chakraborty, "Performance assessment of a bifacial PV system using a new energy estimation model," Solar Energy, vol. 262, p. 111818, 2023.',
'[20] K. Ganesan, D. P. Winston, S. Sugumar and S. Jegan, "Performance analysis of n-type PERT bifacial solar PV module under diverse albedo conditions," Solar Energy, vol. 252, pp. 81-90, 2023.',
'[21] S. A. Pelaez, C. Deline, S. M. MacAlpine, B. Marion, J. S. Stein and R. K. Kostuk, "Comparison of bifacial solar irradiance model predictions with field validation," IEEE Journal of Photovoltaics, vol. 9, no. 1, pp. 82-88, 2019.',
'[22] C. Ghenai, F. F. Ahmad, O. Rejeb and A. K. Hamid, "Sensitivity analysis of design parameters and power gain correlations of bi-facial solar PV system using response surface methodology," Solar Energy, vol. 223, pp. 44-53, 2021.',
]
for ref in refs: P(ref, al=A.LEFT)
doc.add_page_break()

# ===================== APPENDIX A =====================
H("APPENDIX A",16,A.CENTER);H("Software Code Excerpts",14,A.CENTER);NL()
P("This appendix contains the key code excerpts used in the implementation of the bifacial PV analysis framework.")
H("A.1  Irradiance Calculation Function (MATLAB)",12)
P("[Insert MATLAB code for calculate_irradiance function implementing equations (1)-(11)]")
NL()
H("A.2  Frontend Interface Code (HTML/JavaScript)",12)
P("[Insert key excerpts of the frontend interface code showing the parameter input forms and API integration]")
NL()
H("A.3  Backend API Integration (Python/Flask)",12)
P("[Insert key excerpts of the backend code showing NASA POWER API integration and parametric sweep coordination]")
doc.add_page_break()

# ===================== APPENDIX B =====================
H("APPENDIX B",16,A.CENTER);H("Simulink Block Diagrams",14,A.CENTER);NL()
P("This appendix contains the detailed Simulink block diagrams used for the bifacial PV module simulation.")
H("B.1  Complete Simulink Model",12)
P("[Insert full Simulink model screenshot showing all blocks and connections]")
NL()
H("B.2  MATLAB Function Block: calculate_irradiance",12)
P("[Insert screenshot or code listing of the MATLAB function block]")
NL()
H("B.3  PV Array Block Configuration",12)
P("[Insert screenshot of PV Array block parameter settings]")
doc.add_page_break()

# ===================== LIST OF PUBLICATIONS =====================
H("LIST OF PUBLICATIONS",14,A.CENTER);NL()
P("The following publication has been produced as part of this project work:")
NL()
P('[1] A. Upadhyay, P. Chakrabarty, Aman and A. Garg, "Performance analysis of bifacial PV system based on different albedos," accepted for presentation at IC2PCT 2026 (International Conference on Clean and Prospective Computing Technologies), 2026.', al=A.LEFT)
doc.add_page_break()

# ===================== BIO-DATA =====================
H("BIO-DATA",16,A.CENTER);NL()
studs=[("Ashish Upadhyay","2022UEE4521"),("Pratyai Chakrabarty","2022UEE4586"),("Aman","2022UEE4532"),("Aditya Garg","2022UEE4503")]
for name,roll in studs:
    H(f"{name}",12)
    P(f"Roll Number: {roll}")
    P(f"Programme: B.Tech. Electrical Engineering")
    P(f"Department: Electrical Engineering")
    P(f"Institution: Netaji Subhas University of Technology (NSUT), Dwarka, New Delhi - 110078")
    P(f"Batch: 2022-2026")
    NL()

# ===================== SAVE FINAL =====================
doc.save('BTP_Final_Report_Expanded.docx')
print("=== FINAL REPORT SAVED: BTP_Final_Report_Expanded.docx ===")

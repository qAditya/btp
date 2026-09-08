"""Chapter 3 Part B: Simulink Model + Frontend Platform + API + Parametric Sweep"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('ch3_partA.docx')

def H(txt, sz=14, al=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(); p.alignment = al
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(txt); r.bold = True
    r.font.size = Pt(sz); r.font.name = 'Times New Roman'

def P(txt, al=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False, sz=12):
    p = doc.add_paragraph(); p.alignment = al
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run(txt)
    r.font.size = Pt(sz); r.font.name = 'Times New Roman'
    r.bold = bold; r.italic = italic

def NL():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

def fig_ph(num, caption):
    NL()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'[Insert Figure {num} here]'); r.italic = True
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'Fig. 3.{num}. {caption}'); r2.bold = True
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11)
    NL()

# ── 3.3 Simulink Model ───────────────────────────────────────────
H("3.3  Simulink Model Development")
P("To translate the irradiance calculations into electrical output predictions, a MATLAB/Simulink "
  "model was built that replicates the behaviour of the bifacial PV module under a wide range "
  "of environmental and installation conditions. The model accepts GHI, DHI, DNI, albedo, tilt "
  "angle, mounting height, and cell temperature as inputs, and returns the complete I-V and P-V "
  "characteristic curves as outputs.")
P("At the heart of the Simulink schematic is a custom MATLAB Function block named "
  "calculate_irradiance, which implements equations (1) through (11) from Section 3.2. Given "
  "a set of instantaneous meteorological and geometric inputs, this block computes G\u209c\u2092\u209c "
  "and passes it as the effective irradiance G\u1d49\u1da0\u1da0 to the PV Array block. The PV "
  "Array block is configured as a single bifacial module parameterised by its standard-test-condition "
  "short-circuit current (I\u209b\u1d9c), open-circuit voltage (V\u2092\u1d9c), maximum-power-point "
  "current (I\u2098\u209a\u209a) and voltage (V\u2098\u209a\u209a), and the NOCT value of 45\u00b0C.")
P("The I-V characteristic sweep is performed by driving the array terminals with a linearly "
  "ramping controlled voltage source V\u1d3f\u1d43\u2098\u209a. As V\u1d3f\u1d43\u2098\u209a rises "
  "from zero to V\u2092\u1d9c, the PV Array block produces a corresponding current response. "
  "Both the instantaneous terminal voltage (V\u2098\u1d49\u1d43\u209b) and current (I\u2098\u1d49\u1d43\u209b) "
  "are collected through a Bus Creator block and logged to the MATLAB workspace via To Workspace "
  "blocks. A downstream Gain block multiplies the two signals to yield instantaneous power, "
  "providing the P-V curve. Separate Scope blocks display the I-V and P-V curves in real time "
  "during simulation, enabling rapid visual verification of the operating point.")
P("Temperature effects on cell performance are incorporated through the NOCT model. Cell "
  "temperature T\u1d9c\u1d49\u2097\u2097 is calculated from the ambient temperature T\u1d43\u2098\u1d47 "
  "and the effective irradiance using the standard relation, and the resulting temperature "
  "deviation from 25\u00b0C is applied to the PV Array block's internal temperature coefficient "
  "model, capturing the voltage reduction that typically accompanies higher operating "
  "temperatures. This ensures that the simulated peak power values faithfully reflect realistic "
  "field conditions rather than idealised standard test conditions.")

fig_ph(3, "Simulink model for the bifacial PV module showing the irradiance calculation "
         "block, PV Array block, ramp voltage source, and measurement scopes.")

P("The Simulink model is invoked programmatically from within the backend Python service using "
  "the MATLAB Engine API for Python. For each combination of installation parameters in the "
  "parametric sweep, the engine sets the corresponding workspace variables, executes the "
  "simulation, retrieves the logged I-V and P-V data, and extracts the peak power and rear "
  "gain percentage before moving to the next configuration. This tight coupling between the "
  "mathematical formulation, the Simulink electrical model, and the Python orchestration layer "
  "is what enables the fully automated sweep described in Section 3.6.")

# ── 3.4 Frontend Platform ─────────────────────────────────────────
H("3.4  Frontend Platform Architecture")
P("The complete analytical system is structured as a four-layer software stack: a browser-based "
  "user interface, a Python/Flask backend application, an analysis and simulation engine, and "
  "an interactive results dashboard. Each layer has a well-defined responsibility, and data "
  "flow is strictly unidirectional from user input through computation to results display. "
  "The overall architecture is depicted in Fig. 3.4.")

fig_ph(4, "Proposed framework for the optimal configuration of bifacial PV systems "
         "using the MATLAB simulation engine and NASA POWER API.")

H("3.4.1  User Interface Layer", 12)
P("A practitioner interacts with the system entirely through a single-page web application "
  "rendered in a standard desktop browser—no local software installation is required. The "
  "interface presents three collapsible input panels. The first panel accepts the city name, "
  "which is resolved to geographic co-ordinates (latitude, longitude) by a server-side lookup "
  "table, and the analysis date or date range. The second panel defines the parametric sweep "
  "bounds: lower and upper limits for tilt angle (\u03b2) and mounting height (h), plus step "
  "sizes for each variable. The third panel allows the user to choose the ground surface type "
  "from a pre-populated dropdown; the associated albedo value is automatically inserted into "
  "the computation. An optimisation objective selector lets the user choose between 'maximum "
  "cumulative energy' and 'maximum rear-side gain' as the ranking criterion.")
P("Input validation is handled client-side using JavaScript before any network request is "
  "made, providing immediate feedback on out-of-range entries. Once all inputs are confirmed "
  "valid, a single button click triggers an asynchronous HTTP POST to the backend API endpoint, "
  "serialising the input parameters as a JSON payload. The interface then displays a loading "
  "indicator while the backend processes the request.")

H("3.4.2  Backend Application Layer", 12)
P("The backend is implemented as a Flask application served by a Gunicorn WSGI server. Upon "
  "receiving the JSON payload from the browser, the backend performs three sequential tasks. "
  "First, it validates the incoming parameters against predefined bounds and raises descriptive "
  "HTTP error responses for any inconsistency. Second, it constructs the NASA POWER API request "
  "URL, executes the HTTP GET call, and parses the JSON response to extract the hourly irradiance "
  "time series. Third, it calls the analysis engine—described in the following subsection—with "
  "the irradiance data and sweep configuration, waits for the ranked results, and serialises "
  "them as a JSON response to the browser.")
P("Error handling is implemented at three levels: network errors when communicating with the "
  "NASA POWER API return a descriptive message to the user; simulation errors within the MATLAB "
  "engine are caught and logged server-side; and unexpected exceptions trigger a generic error "
  "response that prevents server crashes from propagating to the user interface.")

H("3.4.3  Analysis and Simulation Engine", 12)
P("The analysis engine is a Python module that implements the parametric sweep logic. It "
  "accepts the hourly irradiance arrays and the sweep configuration, iterates over every "
  "combination of tilt angle and mounting height values, and for each combination calls the "
  "MATLAB-based irradiance and electrical models. The MATLAB Engine API for Python launches "
  "a persistent MATLAB session at server start-up; subsequent calls to this session incur "
  "negligible start-up overhead, keeping per-configuration computation times manageable even "
  "for large sweep grids.")
P("For each evaluated configuration, the engine computes: (i) the hourly G\u209c\u2092\u209c "
  "time series using equations (1) through (11); (ii) the daily cumulative energy by numerical "
  "integration of the power time series; (iii) the peak power at the maximum irradiance hour; "
  "and (iv) the rear gain percentage, defined as the fraction of G\u209c\u2092\u209c attributable "
  "to the rear-surface contribution. The results for all configurations are stored in a pandas "
  "DataFrame, sorted by the chosen objective, and returned to the backend layer.")

H("3.4.4  Results Dashboard", 12)
P("The results dashboard renders two complementary views. A sortable HTML table lists the "
  "top-ranked configurations, each row showing the tilt angle, mounting height, surface type, "
  "cumulative energy (kWh), peak power (kW), and rear gain percentage. The first-ranked row "
  "is highlighted to draw attention to the recommended configuration. Alongside the table, "
  "interactive Chart.js plots display the I-V and P-V curves for the top configurations on "
  "a common axis, allowing the user to compare the shape and peak of the characteristic "
  "curves across different parameter settings. A secondary panel provides single-parameter "
  "analysis results—rear gain vs. albedo, rear gain vs. height, and rear gain vs. tilt—which "
  "are generated by fixing two parameters at mid-range values and sweeping the third.")

# ── 3.5 NASA POWER API Integration ──────────────────────────────
H("3.5  NASA POWER API Integration")
P("Meteorological data for the analysis is obtained from the NASA Prediction of Worldwide "
  "Energy Resources (POWER) project, which provides hourly solar radiation and meteorological "
  "parameters derived from the Modern-Era Retrospective analysis for Research and Applications "
  "Version 2 (MERRA-2) reanalysis product and additional satellite-based datasets. The POWER "
  "API delivers data at a spatial resolution of 0.5\u00b0 latitude by 0.625\u00b0 longitude, "
  "which is sufficiently fine for engineering-level analysis of site-specific irradiance.")
P("The backend constructs a POWER API request specifying the latitude and longitude of the "
  "selected city, the analysis date range, and the required parameters: ALLSKY_SFC_SW_DWN "
  "(GHI, W m\u207b\u00b2), ALLSKY_SFC_SW_DIFF (DHI, W m\u207b\u00b2), "
  "ALLSKY_SFC_SW_DNI (DNI, W m\u207b\u00b2), and T2M (ambient air temperature at 2 m, \u00b0C). "
  "The API response is a JSON object containing hourly arrays for each requested parameter. "
  "The backend unpacks this object, aligns the arrays by UTC hour, and converts them to "
  "NumPy arrays for downstream numerical processing.")
P("For the case study described in Chapter 4, the location Greater Noida, India "
  "(lat. 28.47\u00b0 N, lon. 77.50\u00b0 E) was selected and the analysis date was fixed as "
  "13 March 2026. The API returned a total GHI of 6.08 kWh m\u207b\u00b2 and a peak GHI of "
  "866.53 W m\u207b\u00b2 for that day, consistent with typical late-winter irradiance levels "
  "at this latitude.")

# ── 3.6 Parametric Sweep Configuration ──────────────────────────
H("3.6  Parametric Sweep Configuration")
P("The parametric sweep systematically evaluates every point on a multi-dimensional grid of "
  "installation parameter combinations. In the implementation used for the case study, the "
  "following ranges and increments were adopted. Mounting height was varied from 50 cm to "
  "450 cm in steps of 50 cm, giving nine discrete height levels. Tilt angle was varied from "
  "10\u00b0 to 50\u00b0 in steps of 10\u00b0, giving five discrete angles. Ground surface type "
  "was selected by the user from the pre-defined list; for the case study, concrete (\u03c1 = 0.30) "
  "was chosen. This yields a grid of 9 \u00d7 5 = 45 evaluation points per surface type.")
P("At each grid point, the analysis engine executes the following sequence: (i) compute the "
  "shadow bounds from equations (7) and (8) for each hour of the day; (ii) evaluate F\u1d65 "
  "from equation (9); (iii) compute G\u1d3f\u2092\u207f\u209c from equation (5) and G\u1d3f\u1d49\u1d43\u1d3f "
  "from equations (10) and (11); (iv) sum the two contributions per equation (1) to obtain "
  "G\u209c\u2092\u209c; (v) pass G\u209c\u2092\u209c to the Simulink model to extract the I-V "
  "and P-V curves; (vi) record cumulative energy, peak power, and rear gain. The 45 result "
  "records are sorted by the chosen objective, and the top five configurations are presented "
  "to the user in the dashboard table.")
P("Beyond the full three-parameter sweep, the platform also supports single-parameter variation "
  "analyses in which two of the three parameters are held fixed and the third is swept over a "
  "finer grid. This mode—accessible from a dedicated section of the user interface—allows "
  "designers to isolate and quantify the individual contribution of each installation variable. "
  "Results from such focused analyses for the case study are presented in Section 4.3.")
P("The azimuth angle of the PV module was kept fixed at 180\u00b0 (south-facing) throughout all "
  "simulations, as this orientation is standard practice for maximising annual solar exposure "
  "in the northern hemisphere. The NOCT was set to 45\u00b0C and the bifaciality factor to 0.70 "
  "for all evaluations.")
NL()

doc.save('ch3_partB.docx')
print("Part B saved.")

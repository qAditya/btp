from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document('BTP_Report_Part1.docx')
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)

def H(text, sz=14, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text); r.bold=True; r.font.size=Pt(sz); r.font.name='Times New Roman'
    return p

def P(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.size=Pt(12); r.font.name='Times New Roman'; r.bold=bold; r.italic=italic
    return p

def NL(n=1):
    for _ in range(n):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)

# ===== CHAPTER 1 =====
H("CHAPTER 1", 16, WD_ALIGN_PARAGRAPH.CENTER)
H("INTRODUCTION", 14, WD_ALIGN_PARAGRAPH.CENTER)
NL()

H("1.1  Background and Motivation")
P("The global energy landscape is undergoing a significant transformation driven by the urgent need to address climate change, reduce greenhouse gas emissions, and transition toward sustainable energy sources. Among the various renewable energy technologies available today, solar photovoltaic (PV) technology has emerged as one of the most promising and rapidly growing solutions for clean electricity generation. The International Energy Agency (IEA) has reported that solar PV is the fastest-growing source of electricity generation worldwide, with global installed capacity exceeding 1,500 GW by the end of 2024. This unprecedented growth is attributed to declining module costs, supportive government policies, and increasing awareness of the environmental impact of fossil fuel-based power generation.")

P("India, in particular, has witnessed remarkable progress in solar energy deployment. With an ambitious target of 500 GW of non-fossil fuel electricity capacity by 2030, the country has emerged as one of the leading markets for solar PV installations globally. India receives approximately 5,000 trillion kWh of solar energy annually, with most parts of the country experiencing 300 or more sunny days per year, making it an ideal location for large-scale solar power deployment. The National Solar Mission and various state-level policies have accelerated the adoption of solar PV systems across residential, commercial, and utility-scale segments.")

P("Within the realm of PV technology, a critical distinction exists between conventional monofacial modules and the more advanced bifacial photovoltaic modules. Conventional monofacial PV panels generate electricity exclusively from solar radiation incident on their front surface, thereby leaving a substantial portion of the available solar resource underutilized. Specifically, diffuse radiation from the sky hemisphere and reflected radiation from the ground and surrounding surfaces remain largely unharnessed in monofacial configurations. Bifacial PV modules, in contrast, are engineered to capture and convert solar irradiance incident on both the front and rear surfaces of the module, thereby achieving higher energy yields from the same installation footprint [1], [2].")

P("The ability of bifacial modules to harvest rear-side irradiance represents a significant technological advancement in solar energy utilization. Several studies have demonstrated that bifacial PV systems can achieve energy yield improvements ranging from 5% to 30% compared to equivalent monofacial installations, depending on the installation configuration and environmental conditions [3], [4]. This additional energy generation potential has made bifacial modules an increasingly attractive option for solar power plant developers and system designers seeking to maximize the return on investment for their installations.")

P("However, realizing the full performance potential of bifacial PV systems is not straightforward. The rear-side irradiance contribution is highly sensitive to a complex interplay of installation parameters and environmental factors. Key parameters including the module tilt angle, ground surface reflectivity (albedo), and mounting height above the ground collectively determine the amount of reflected and diffuse radiation reaching the rear surface of the module [5], [6]. Improper selection of these parameters can lead to suboptimal system performance, negating the inherent advantages of bifacial technology.")

P("This underscores the critical need for systematic optimization tools and frameworks that can evaluate the combined influence of multiple installation parameters on bifacial PV system performance for specific geographical locations. The present work addresses this need by developing a comprehensive computational framework for parametric analysis and optimization of bifacial PV installations.")

H("1.2  Solar Photovoltaic Technology: An Overview")
P("Solar photovoltaic technology converts sunlight directly into electricity using semiconductor materials that exhibit the photovoltaic effect. When photons from sunlight strike the surface of a PV cell, they transfer their energy to electrons in the semiconductor material, creating electron-hole pairs. These charge carriers are separated by the built-in electric field of the p-n junction within the cell, generating a direct current (DC) that can be collected and utilized.")

P("Modern PV cells are predominantly based on crystalline silicon technology, which accounts for approximately 95% of the global PV module production. Crystalline silicon cells are further categorized into monocrystalline and polycrystalline variants, with monocrystalline cells offering higher conversion efficiencies typically ranging from 20% to 24% under standard test conditions. Recent advancements in cell architectures, such as Passivated Emitter and Rear Contact (PERC), Tunnel Oxide Passivated Contact (TOPCon), and Heterojunction Technology (HJT), have pushed the boundaries of achievable efficiencies while simultaneously reducing manufacturing costs.")

P("The electrical output of a PV module is characterized by its current-voltage (I-V) relationship, which describes the module's behavior under specific irradiance and temperature conditions. The maximum power point (MPP) on the I-V curve represents the operating condition at which the module delivers maximum electrical power. The power-voltage (P-V) characteristic curve provides a direct visualization of the power output as a function of operating voltage, with the peak of this curve corresponding to the MPP.")

H("1.3  Monofacial versus Bifacial PV Modules")
P("Conventional monofacial PV modules feature an opaque rear surface, typically covered with a white or colored backsheet, which prevents any light transmission to the rear side of the cells. As a result, only the radiation incident on the front surface of the module contributes to electricity generation. While this configuration has served the solar industry well for several decades, it inherently limits the module's ability to capture all available solar radiation at the installation site.")

P("Bifacial PV modules, in contrast, feature transparent or semi-transparent rear surfaces that allow light to reach the rear side of the solar cells. These modules utilize solar cells with metallization patterns on both the front and rear surfaces, enabling photocurrent generation from radiation incident on either side. The bifaciality factor, defined as the ratio of rear-side efficiency to front-side efficiency under standard test conditions, is a key parameter characterizing the performance capability of bifacial modules. Modern bifacial modules typically exhibit bifaciality factors ranging from 0.70 to 0.90, indicating that the rear side can generate 70% to 90% of the power produced by the front side under equivalent irradiance conditions [1].")

P("The additional energy generated by the rear surface of bifacial modules is commonly quantified using the concept of bifacial gain, which represents the percentage increase in energy yield compared to an equivalent monofacial installation. Research has shown that bifacial gains can range from as low as 5% in poorly optimized installations to over 30% in configurations with highly reflective ground surfaces and elevated mounting heights [4], [11].")

H("1.4  Key Parameters Affecting Bifacial PV Performance")
P("The performance of bifacial PV systems is governed by a complex interaction between solar radiation characteristics, installation geometry, and ground surface properties. The primary parameters affecting the rear-side irradiance contribution are discussed below.")

P("Tilt Angle: The module tilt angle significantly influences the balance between front-side and rear-side irradiance capture. For front-side performance, the optimal tilt angle is primarily determined by the latitude of the installation site, as it maximizes the annual direct beam irradiance on the module surface. However, for bifacial modules, the tilt angle also affects the sky view factor and ground view factor, which determine the proportions of diffuse sky radiation and ground-reflected radiation reaching the module surfaces. Studies have shown that the optimal tilt angle for bifacial modules may differ from that of monofacial modules due to the additional consideration of rear-side irradiance maximization [5], [6].", bold=False, italic=False)

P("Ground Albedo: The ground surface reflectivity, commonly referred to as albedo, plays a crucial role in determining the amount of radiation reflected toward the rear surface of the bifacial module. Albedo values vary significantly across different surface types, ranging from approximately 0.08 for very dirty galvanized surfaces to 0.85 for aluminum surfaces. High-albedo surfaces such as snow, white-painted concrete, and metallic surfaces can substantially increase rear-side irradiance and consequently improve the overall energy generation of bifacial PV systems [7], [8], [9], [10].")

P("Mounting Height: The elevation of PV modules above the ground affects the geometric view factor between the module and the reflective ground surface. Higher mounting positions generally allow the rear surface of the module to 'see' a larger area of the ground, thereby increasing the amount of reflected radiation received. However, the incremental benefit of additional height diminishes beyond certain elevations, and practical considerations such as structural costs and wind loading must be balanced against the marginal energy gains [11].")

P("Shadow and View Factor: In array configurations, the shadow cast by each module on the ground behind it reduces the effective ground area from which reflected radiation can reach the rear surface. The shadow view factor quantifies this reduction and is determined by the solar position, module geometry, tilt angle, and mounting height. Accurate modeling of shadow effects is essential for realistic estimation of rear-side irradiance in multi-row installations [14], [18].")

H("1.5  Challenges in Bifacial PV Optimization")
P("Despite the growing body of research on bifacial photovoltaic systems, several challenges remain in determining the optimal configuration of these systems under varying environmental conditions. Many existing studies have focused on individual parameters such as tilt angle, albedo, or module elevation, often evaluating their effects independently rather than considering their combined influence on system performance [5], [8], [9]. This approach provides limited insight into the complex interdependencies between these parameters.")

P("Other investigations have developed modeling approaches or global performance assessments of bifacial PV systems; however, these studies are often limited to predefined system configurations or specific climatic conditions [4]. Furthermore, research on the design of large-scale bifacial solar farms indicates that multiple parameters, including module orientation, mounting height, and system geometry, must be optimized simultaneously to achieve maximum energy output [14]. The lack of accessible, user-friendly tools for location-specific parametric optimization of bifacial PV systems represents a significant gap in the current state of the art.")

P("Additionally, existing optimization approaches often require specialized software knowledge, access to commercial simulation tools, or significant computational resources, making them impractical for many PV system designers and installers who need quick, reliable optimization guidance for specific project sites.")

H("1.6  Problem Statement and Research Objectives")
P("Based on the identified gaps in the existing literature, this project addresses the need for a systematic analytical framework that can evaluate the combined influence of key installation parameters using location-specific solar irradiance data. The primary objectives of this work are as follows:")

P("1. To develop a comprehensive mathematical model for bifacial PV system performance that accounts for front-side and rear-side irradiance contributions, including the effects of tilt angle, ground albedo, mounting height, and view factor geometry.")
P("2. To implement the mathematical model in a MATLAB/Simulink environment for detailed electrical simulation of bifacial PV module characteristics under varying installation configurations.")
P("3. To design and develop a user-oriented frontend platform that integrates location-based solar irradiance data retrieval from the NASA POWER API with parametric performance analysis.")
P("4. To perform parametric sweep analyses to identify optimal installation configurations that maximize energy generation or rear-side gain for a given geographical location.")
P("5. To validate the proposed framework through a case study and compare the results with published literature.")

H("1.7  Organization of the Report")
P("The remainder of this report is organized as follows. Chapter 2 presents a comprehensive review of the existing literature on bifacial PV technology, covering studies on the influence of tilt angle, ground albedo, mounting height, and combined parametric optimization. Chapter 3 describes the methodology adopted in this work, including the mathematical formulation of bifacial PV operation, the Simulink model development, and the frontend platform architecture. Chapter 4 presents the results of the parametric analyses conducted for the case study location and provides a detailed discussion of the findings. Chapter 5 concludes the report by summarizing the key contributions, identifying the limitations of the present work, and suggesting directions for future research.")

doc.add_page_break()

# ===== CHAPTER 2 =====
H("CHAPTER 2", 16, WD_ALIGN_PARAGRAPH.CENTER)
H("LITERATURE REVIEW", 14, WD_ALIGN_PARAGRAPH.CENTER)
NL()

H("2.1  Review of Bifacial PV Technology")
P("Bifacial photovoltaic technology has undergone significant evolution since its conceptual introduction in the 1960s and the first industrial production of bifacial PV panels in 1984 [7 in btp1]. The technology has gained increasing commercial traction in recent years, with market projections indicating that bifacial modules could capture approximately 35% of the global PV market by 2027. This growing market share is driven by the inherent ability of bifacial modules to generate additional electricity from the rear surface without requiring a proportional increase in installation area.")

P("Guerrero-Lemus et al. [1] provided a comprehensive review of bifacial solar photovoltaic technology, covering the historical development, cell architectures, and performance characteristics of bifacial modules. Their review highlighted that rear-side irradiance contribution is strongly influenced by environmental and geometric parameters including ground reflectivity, module tilt angle, and installation height. The study also identified the need for standardized testing and characterization protocols for bifacial modules.")

P("Alam et al. [2] conducted a performance comparison between bifacial and monofacial photovoltaic systems under varying environmental conditions. Their investigation demonstrated that bifacial systems consistently outperformed monofacial counterparts across diverse climatic scenarios, with the magnitude of improvement dependent on the local irradiance conditions and installation configuration. The study reinforced the importance of site-specific analysis for quantifying the expected benefits of bifacial technology.")

P("Deline et al. [3] from the National Renewable Energy Laboratory (NREL) presented an important study on separating fact from fiction in bifacial PV system performance. Their work provided empirical evidence to validate the performance claims associated with bifacial technology and established benchmarks for expected energy yield improvements under various installation conditions. The NREL study has become a foundational reference for subsequent research on bifacial PV systems.")

P("Sun et al. [4] undertook a global perspective analysis of the optimization and performance of bifacial solar modules. Their study demonstrated that bifacial modules offer significant energy yield improvements across diverse geographic locations when properly configured. The research highlighted the importance of considering location-specific solar resource characteristics and installation parameters in optimizing bifacial PV system design.")

H("2.2  Influence of Tilt Angle on Bifacial PV Performance")
P("The tilt angle of a PV module is a fundamental design parameter that determines the angle of incidence of solar radiation on the module surface. For monofacial modules, the optimal tilt angle is primarily governed by the latitude of the installation site, with adjustments for local climatic conditions and the desired seasonal energy distribution.")

P("Basak et al. [5] investigated tilt angle optimization specifically for bifacial photovoltaic modules, demonstrating that the optimal tilt angle for bifacial modules may differ from that of monofacial modules due to the additional contribution of ground-reflected radiation to the rear surface. Their study showed that increasing the tilt angle beyond the latitude-optimal value can enhance rear-side performance, although this may come at the cost of reduced front-side irradiance capture.")

P("Peter and Novak [6] examined the optimal tilt angle for maximizing energy production of bifacial solar panels. Their findings indicated that the interplay between front-side and rear-side irradiance components creates a more complex optimization landscape for bifacial modules compared to monofacial systems. The study recommended that tilt angle optimization for bifacial installations should account for both direct and reflected radiation components.")

P("Yusufoglu et al. [18] presented a detailed simulation study examining the annual energy yield of south-facing bifacial modules with rigorous calculation of ground-reflected radiation. Their work demonstrated that optimum tilt angles of bifacial modules depend on a larger set of parameters than standard modules, requiring the albedo coefficient and installation height to be taken into account. They found that at low module elevations, the distance between shadows and modules is shorter, and this drawback can be reduced by increasing tilt angles. They showed that the difference in optimum tilt angle for standard and bifacial modules can vary by 3 to 5 degrees depending on the location and albedo coefficient. This seminal work provided key insights into the view factor calculation methodology adopted in the present study.")

H("2.3  Effect of Ground Albedo on Rear-Side Irradiance")
P("Ground surface albedo is widely recognized as the most influential parameter affecting the rear-side energy contribution of bifacial PV systems. The albedo coefficient quantifies the fraction of incident solar radiation that is reflected by a surface, with values ranging from near zero for highly absorptive surfaces to values approaching unity for highly reflective surfaces.")

P("Riedel-Lyngskær et al. [7] investigated the effect of spectral albedo on the energy yield of bifacial photovoltaic systems. Their study highlighted that the spectral characteristics of the reflected radiation, in addition to the overall albedo magnitude, can influence the electrical output of bifacial modules due to the wavelength-dependent response of silicon solar cells.")

P("Aksoy and Ceylan [8] conducted an investigation of ground albedo effects on bifacial photovoltaic panel performance. Their research provided quantitative data on the relationship between ground surface type and rear-side energy generation, demonstrating that surface selection is a critical design consideration for bifacial PV installations.")

P("Baghel et al. [9] performed a comprehensive performance evaluation and optimization study examining both albedo and tilt angle in bifacial photovoltaic systems. Their work demonstrated the synergistic effects of these two parameters and emphasized the need for simultaneous optimization to achieve maximum system performance.")

P("Atalay et al. [10] conducted experimental investigations of ground surface reflectivity impact on bifacial photovoltaic systems. Their field measurements provided valuable empirical data on the actual albedo values achievable with various ground surface types under realistic operating conditions, complementing the theoretical and simulation-based studies in the literature.")

P("Dincer and Ozer [11] presented a comprehensive parametric analysis investigating the optimization of rear-side energy contribution in bifacial PV panels. Their study systematically evaluated four key parameters: albedo, tilt angle, panel height, and mounting configuration. Using experimentally measured GHI data from Turkey, they found that aluminum surfaces provided the highest rear-surface energy generation of 21.2%, followed by fresh snow at 20.5%. Their analysis demonstrated that increasing the tilt angle from 10 degrees to 50 degrees led to a gradual rise in rear-side gain from 4.3% to 5.5%. For panel height variation, the gain increased from 4.1% at 40 cm to 4.5% at 100 cm. They also compared horizontal and vertical mounting configurations, finding that horizontal mounting provided a slightly higher rear-side energy yield of 4.5% compared to 4.1% for vertical mounting. This study serves as a primary reference for the parametric analysis approach adopted in the present work.")

H("2.4  Impact of Mounting Height on Energy Yield")
P("The height at which bifacial PV modules are mounted above the ground surface directly affects the geometric relationship between the module and the reflective ground, which in turn influences the amount of ground-reflected radiation reaching the rear surface of the module.")

P("Dincer and Ozer [11] investigated the effect of panel height on rear irradiation performance, evaluating heights from 40 cm to 100 cm. Their results indicated a gradual increase in rear-side energy generation with increasing panel height, though the differences between successive heights were relatively small, suggesting diminishing returns at higher elevations.")

P("Yusufoglu et al. [18] provided important insights into the influence of module elevation on bifacial PV performance through their simulation study. They demonstrated that the annual energy yield exhibits a near-logarithmic dependence on module elevation, with the benefit being more pronounced for locations with more direct radiation (such as Cairo) compared to locations with predominantly diffuse radiation (such as Oslo). Their simulations showed that at optimum tilt angles, the produced annual energy can be increased by up to 30% compared to a standard module simply by positioning modules two meters above the ground instead of close-to-ground installation. Furthermore, they demonstrated a linear relationship between albedo coefficient and annual energy yield for all module elevation heights.")

P("Braga et al. [12] examined the performance of bifacial PV modules under tropical climatic conditions, where high solar elevation angles and intense direct radiation create unique challenges for rear-side irradiance optimization. Their study provided valuable data on the interaction between mounting height and tropical solar geometry.")

P("Almarshoud et al. [13] conducted experimental performance analysis of bifacial photovoltaic modules under different operating conditions, including variations in mounting height. Their field data confirmed the positive correlation between module elevation and rear-side energy contribution, while also highlighting the practical constraints that limit the feasible range of mounting heights in real installations.")

H("2.5  Combined Parametric Optimization Studies")
P("While many studies have examined individual parameters in isolation, a growing body of research has recognized the importance of evaluating the combined effects of multiple parameters on bifacial PV system performance.")

P("Patel et al. [14] addressed the optimum design of tracking bifacial solar farms, demonstrating that multiple parameters including module orientation, mounting height, and system geometry must be optimized simultaneously to achieve maximum energy output. Their work highlighted the limitations of single-parameter optimization approaches and advocated for comprehensive parametric analysis frameworks.")

P("Sun et al. [4] provided a global perspective on bifacial module optimization, evaluating the combined effects of tilt angle, albedo, and latitude on system performance across diverse geographic locations. Their findings demonstrated that the relative importance of different installation parameters varies significantly with geographic location, underscoring the need for location-specific optimization tools.")

H("2.6  Simulation and Modeling Approaches")
P("Accurate simulation and modeling of bifacial PV systems require the integration of solar radiation models, geometric view factor calculations, and electrical performance models. Several approaches have been developed and validated in the literature.")

P("Liu and Jordan [15] developed the isotropic sky model, which forms the foundation for estimating tilted surface irradiance from horizontal irradiance measurements. This model decomposes the global horizontal irradiance into beam, diffuse, and ground-reflected components, providing a practical framework for estimating the irradiance on tilted PV surfaces.")

P("Beckman and Duffie [16] provided comprehensive analytical methods for solar engineering applications, including detailed formulations for solar geometry calculations, irradiance decomposition, and tilted surface irradiance estimation. Their work remains a standard reference for solar energy system design and analysis.")

P("Marion et al. [17] developed a practical irradiance model specifically for bifacial PV modules, accounting for the unique characteristics of rear-side irradiance including ground-reflected and diffuse components. Their model addresses the shadow effects of module rows on the ground and provides an improved estimation of the available rear-side irradiance.")

P("Yusufoglu et al. [18] demonstrated an approach accounting for the shadow effect on ground-reflected irradiance based on the principle of the view factor known from heat transfer fundamentals. Their model separates the GHI into diffuse and direct components, recognizing that the diffuse component remains unchanged while the direct portion of ground-reflected radiation is emitted only from the area outside the shadow. This methodology forms the basis of the rear-side irradiance model implemented in the present work.")

H("2.7  Research Gap Analysis")
P("Based on the comprehensive review of existing literature, several key research gaps have been identified that motivate the present work:")

P("1. Limited Combined Analysis: Most existing studies evaluate installation parameters independently rather than considering their combined influence on bifacial PV system performance. There is a need for frameworks that can simultaneously analyze the effects of tilt angle, albedo, and mounting height.")
P("2. Location-Specific Optimization: Many studies are limited to specific geographic locations or generalized climatic conditions. A flexible framework that can retrieve and utilize location-specific irradiance data for any given site is needed.")
P("3. Accessibility Gap: Current optimization approaches often require specialized software knowledge or access to commercial simulation tools, creating a barrier for many PV system designers.")
P("4. Dual Optimization Objectives: Few studies consider the trade-off between maximum total energy generation and maximum rear-side gain as separate but complementary optimization objectives.")

P("The present work addresses these gaps by developing a comprehensive computational framework that integrates location-specific solar irradiance data retrieval, rigorous mathematical modeling of bifacial PV operation, and a user-friendly frontend platform for parametric analysis and optimization.")

# Save table of literature
H("Table 2.1: Albedo Values of Common Surface Types for PV Applications [11]", sz=12)
tbl = doc.add_table(rows=14, cols=2)
tbl.style = 'Table Grid'
headers = [("Surface Type", "Albedo Value")]
data = [
    ("Urban environment", "0.14 - 0.22"),
    ("Grass", "0.15 - 0.25"),
    ("Fresh grass", "0.26"),
    ("Fresh snow", "0.82"),
    ("Wet snow", "0.55 - 0.75"),
    ("Dry asphalt", "0.09 - 0.15"),
    ("Wet asphalt", "0.18"),
    ("Concrete", "0.25 - 0.35"),
    ("Red tiles", "0.33"),
    ("Aluminum", "0.85"),
    ("Copper", "0.74"),
    ("New galvanized steel", "0.35"),
    ("Very dirty galvanized", "0.08"),
]
for j, h in enumerate(headers[0]):
    cell = tbl.cell(0, j); cell.text = h
    for p in cell.paragraphs:
        for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11)
for i, (s, a) in enumerate(data):
    tbl.cell(i+1, 0).text = s
    tbl.cell(i+1, 1).text = a
    for j in range(2):
        for p in tbl.cell(i+1, j).paragraphs:
            for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(11)

doc.add_page_break()
doc.save('BTP_Report_Ch1Ch2.docx')
print("Chapters 1-2 saved!")

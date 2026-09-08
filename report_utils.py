"""Utility functions for report generation with proper equation support."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

def setup_doc(doc=None):
    if doc is None:
        doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1); s.bottom_margin = Inches(1)
        s.left_margin = Inches(1.5); s.right_margin = Inches(1)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'; style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    return doc

def H(doc, text, sz=14, al=WD_ALIGN_PARAGRAPH.LEFT, bold=True):
    p = doc.add_paragraph(); p.alignment = al
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(sz); r.font.name = 'Times New Roman'
    return p

def P(doc, text, bold=False, italic=False, al=WD_ALIGN_PARAGRAPH.JUSTIFY, sz=12):
    p = doc.add_paragraph(); p.alignment = al
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run(text); r.font.size = Pt(sz); r.font.name = 'Times New Roman'
    r.bold = bold; r.italic = italic
    return p

def NL(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

def _mr(oMath, text, italic=True, font='Cambria Math'):
    """Add a math run to an oMath element."""
    r = OxmlElement('m:r')
    if not italic:
        rPr = OxmlElement('m:rPr')
        sty = OxmlElement('m:sty'); sty.set(qn('m:val'), 'p')
        rPr.append(sty); r.append(rPr)
    # Word run properties for font
    wRPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    wRPr.append(rFonts)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '24')
    wRPr.append(sz)
    r.append(wRPr)
    t = OxmlElement('m:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    oMath.append(r)
    return r

def _frac(oMath, num_text, den_text):
    """Add a fraction to an oMath element."""
    f = OxmlElement('m:f')
    fPr = OxmlElement('m:fPr')
    f.append(fPr)
    num = OxmlElement('m:num')
    _mr(num, num_text)
    f.append(num)
    den = OxmlElement('m:den')
    _mr(den, den_text)
    f.append(den)
    oMath.append(f)

def _sub(oMath, base, subscript):
    """Add subscript to oMath."""
    sSub = OxmlElement('m:sSub')
    e = OxmlElement('m:e'); _mr(e, base)
    sub = OxmlElement('m:sub'); _mr(sub, subscript)
    sSub.append(e); sSub.append(sub)
    oMath.append(sSub)

def _sup(oMath, base, superscript):
    """Add superscript to oMath."""
    sSup = OxmlElement('m:sSup')
    e = OxmlElement('m:e'); _mr(e, base)
    sup = OxmlElement('m:sup'); _mr(sup, superscript)
    sSup.append(e); sSup.append(sup)
    oMath.append(sSup)

def add_equation(doc, eq_num, build_fn):
    """Add a centered equation with number on the right.
    build_fn(oMath) should populate the oMath element.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Add oMath element
    oMathPara = OxmlElement('m:oMathPara')
    oMathParaPr = OxmlElement('m:oMathParaPr')
    jc = OxmlElement('m:jc'); jc.set(qn('m:val'), 'center')
    oMathParaPr.append(jc)
    oMathPara.append(oMathParaPr)
    
    oMath = OxmlElement('m:oMath')
    build_fn(oMath)
    oMathPara.append(oMath)
    p._element.append(oMathPara)
    
    # Add equation number as right-aligned tab
    run = p.add_run(f'\t({eq_num})')
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    return p

def add_simple_equation(doc, equation_text, eq_num):
    """Add equation using Cambria Math font, centered, with number."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    # Use OMML
    oMathPara = OxmlElement('m:oMathPara')
    oMath = OxmlElement('m:oMath')
    _mr(oMath, equation_text)
    oMathPara.append(oMath)
    p._element.append(oMathPara)
    
    # Equation number
    run = p.add_run(f'          ({eq_num})')
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    return p

def add_table_grid(doc, headers, rows, caption=None, cap_above=True):
    """Add a formatted table with optional caption."""
    if caption and cap_above:
        H(doc, caption, 12, WD_ALIGN_PARAGRAPH.CENTER)
    
    tbl = doc.add_table(rows=len(rows)+1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for j, h in enumerate(headers):
        c = tbl.cell(0, j); c.text = ''
        cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(h); r.bold = True
        r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = tbl.cell(i+1, j); c.text = ''
            cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cp.add_run(str(v))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    
    if caption and not cap_above:
        P(doc, caption, bold=True, al=WD_ALIGN_PARAGRAPH.CENTER, sz=11)
    return tbl

def fig_placeholder(doc, fig_num, caption):
    """Add a figure placeholder."""
    NL(doc)
    P(doc, f'[Insert {fig_num} here]', al=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    P(doc, f'{fig_num}: {caption}', bold=True, al=WD_ALIGN_PARAGRAPH.CENTER, sz=11)
    NL(doc)

print("Utils module ready.")

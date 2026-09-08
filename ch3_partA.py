"""Chapter 3 Part A: Overview + Mathematical Formulation (Eqs 1-11)"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin    = Inches(1);   s.bottom_margin = Inches(1)
    s.left_margin   = Inches(1.5); s.right_margin  = Inches(1)
ns = doc.styles['Normal']
ns.font.name = 'Times New Roman'; ns.font.size = Pt(12)

def H(txt, sz=14, al=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(); p.alignment = al
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(txt); r.bold = True
    r.font.size = Pt(sz); r.font.name = 'Times New Roman'

def P(txt, al=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(); p.alignment = al
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run(txt)
    r.font.size = Pt(12); r.font.name = 'Times New Roman'

def NL():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

def _mr(parent, txt, italic=True):
    r = OxmlElement('m:r')
    if not italic:
        rPr = OxmlElement('m:rPr')
        sty  = OxmlElement('m:sty'); sty.set(qn('m:val'), 'p')
        rPr.append(sty); r.append(rPr)
    wRPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Cambria Math'); rf.set(qn('w:hAnsi'), 'Cambria Math')
    wRPr.append(rf); r.append(wRPr)
    t = OxmlElement('m:t'); t.text = txt
    t.set(qn('xml:space'), 'preserve'); r.append(t)
    parent.append(r)

def _frac(parent, n, d):
    f = OxmlElement('m:f')
    fPr = OxmlElement('m:fPr'); f.append(fPr)
    nu = OxmlElement('m:num'); _mr(nu, n); f.append(nu)
    de = OxmlElement('m:den'); _mr(de, d); f.append(de)
    parent.append(f)

def _sub(parent, base, sub):
    s = OxmlElement('m:sSub')
    e = OxmlElement('m:e'); _mr(e, base); s.append(e)
    sb = OxmlElement('m:sub'); _mr(sb, sub); s.append(sb)
    parent.append(s)

def _ssub(parent, base, sub):
    """subscript where base is already an element"""
    s = OxmlElement('m:sSub'); s.append(base)
    sb = OxmlElement('m:sub'); _mr(sb, sub); s.append(sb)
    parent.append(s)

def add_eq(build_fn, num):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    oMP  = OxmlElement('m:oMathPara')
    oPr  = OxmlElement('m:oMathParaPr')
    jc   = OxmlElement('m:jc'); jc.set(qn('m:val'), 'center')
    oPr.append(jc); oMP.append(oPr)
    oM = OxmlElement('m:oMath')
    build_fn(oM)
    oMP.append(oM); p._element.append(oMP)
    run = p.add_run(f'          ({num})')
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

# ── CHAPTER HEADING ──────────────────────────────────────────────
H("CHAPTER 3", 16, WD_ALIGN_PARAGRAPH.CENTER)
H("METHODOLOGY", 14, WD_ALIGN_PARAGRAPH.CENTER)
NL()

# ── 3.1 Overview ─────────────────────────────────────────────────
H("3.1  Overview of the Proposed Framework")
P("This chapter details the methodology developed for the parametric analysis and performance "
  "optimisation of bifacial photovoltaic systems. The framework rests on three interlocking "
  "pillars: a physics-based irradiance model that resolves both front- and rear-surface "
  "contributions; a MATLAB/Simulink electrical model that translates the computed irradiance "
  "into current-voltage and power-voltage characteristics; and a web-based frontend platform "
  "that automates data retrieval and parametric sweep execution for any user-specified location "
  "and date range. Together, these components form an end-to-end pipeline through which a "
  "practitioner can determine the installation geometry—tilt angle, ground albedo, and mounting "
  "height—that maximises energy harvest at a chosen site without requiring proprietary software "
  "or specialist modelling knowledge.")
P("The computational sequence begins when the user specifies a geographic location, date range, "
  "and sweep bounds through the browser interface. The backend service converts the city name to "
  "geographic co-ordinates, fetches hourly irradiance data from the NASA POWER API, and feeds "
  "this dataset into the analysis engine. The engine evaluates every combination of the swept "
  "parameters using the mathematical formulation described in Sections 3.2 and 3.3, ranks the "
  "resulting configurations according to the chosen objective—maximum cumulative energy or "
  "maximum rear-side gain—and passes the ranked table together with I-V/P-V curves to the "
  "results dashboard. The overall framework architecture is illustrated in Fig. 3.4.")

# ── 3.2 Mathematical Formulation ─────────────────────────────────
H("3.2  Mathematical Formulation of Bifacial PV Operation")
P("A monofacial module converts only the radiation striking its glazed front surface. Bifacial "
  "modules, by contrast, employ transparent or semi-transparent rear encapsulants and bifacially "
  "active cells so that ground-reflected and diffuse radiation reaching the rear surface also "
  "contribute to power generation [1], [4]. The total effective irradiance received by the "
  "module is accordingly the sum of the two surface contributions:")

def eq1(oM):
    _sub(oM, 'G', 'tot'); _mr(oM, '  =  ')
    _sub(oM, 'G', 'front'); _mr(oM, '  +  '); _sub(oM, 'G', 'rear')
add_eq(eq1, 1)

P("where G\u209c\u2092\u209c is the total effective irradiance incident on the bifacial module "
  "(W m\u207b\u00b2), G\u1da0\u1d63\u2092\u207f\u209c is the irradiance received on the front "
  "surface, and G\u1d63\u1d49\u1d43\u1d3f is the irradiance received on the rear surface. "
  "Both components are functions of solar position, module geometry, and ground reflectivity, "
  "and are determined through the sub-models presented below.")

# ── 3.2.1 Front-Side ─────────────────────────────────────────────
H("3.2.1  Front-Side Irradiance Model", 12)
P("Front-surface irradiance is resolved using the Liu–Jordan isotropic sky model [15], which "
  "decomposes global horizontal irradiance (GHI) into three physically distinct components: "
  "the direct beam, the isotropic diffuse sky, and diffuse ground reflection. Each component "
  "responds differently to module orientation and must therefore be treated separately.")
P("The beam component incident on the tilted surface is found by projecting the direct normal "
  "irradiance (DNI) onto the module plane. A floor function prevents negative contributions "
  "when the sun dips below the effective module horizon:")

def eq2(oM):
    _mr(oM, 'beam'); _sub(oM, '', 'tilted'); _mr(oM, '  =  DNI')
    _mr(oM, ' \u00d7 max(0, cos '); _mr(oM, 'I', italic=True); _mr(oM, ')')
add_eq(eq2, 2)

P("Here DNI is the direct normal irradiance (W m\u207b\u00b2) and I is the angle of incidence "
  "between the incoming beam and the module surface normal. The max(0, cos I) clipping "
  "eliminates physically unrealisable negative values when I > 90\u00b0.")
P("For a tilted surface, the fraction of the isotropic sky dome visible from above is "
  "characterised by the sky view factor, while the fraction of the ground visible from below "
  "defines the ground view factor [15]. Their expressions are:")

def eq3(oM):
    _mr(oM, 'sky'); _sub(oM, '', 'vf'); _mr(oM, '  =  ')
    _frac(oM, '1 + cos\u03b2', '2')
add_eq(eq3, 3)

def eq4(oM):
    _mr(oM, 'gnd'); _sub(oM, '', 'vf'); _mr(oM, '  =  ')
    _frac(oM, '1 \u2212 cos\u03b2', '2')
add_eq(eq4, 4)

P("where \u03b2 is the module tilt angle measured from the horizontal plane. As \u03b2 increases "
  "toward vertical, sky\u1d65\u1da0 decreases and gnd\u1d65\u1da0 increases, which shifts the "
  "balance between sky-diffuse and ground-reflected contributions to the front surface. "
  "Combining the three components, the total front-surface irradiance is [15]:")

def eq5(oM):
    _sub(oM, 'G', 'front'); _mr(oM, '  =  beam')
    _sub(oM, '', 'tilted'); _mr(oM, '  +  (DHI \u00d7 sky')
    _sub(oM, '', 'vf'); _mr(oM, ')  +  (GHI \u00d7 \u03c1 \u00d7 gnd')
    _sub(oM, '', 'vf'); _mr(oM, ')')
add_eq(eq5, 5)

P("DHI (W m\u207b\u00b2) is the diffuse horizontal irradiance, GHI (W m\u207b\u00b2) is the "
  "global horizontal irradiance, and \u03c1 is the dimensionless ground surface albedo "
  "coefficient whose values for common installation surfaces are listed in Table 2.1.")

# ── 3.2.2 Rear-Side ──────────────────────────────────────────────
H("3.2.2  Rear-Side Irradiance Model", 12)
P("Computing rear irradiance is more involved than the front-side calculation because the "
  "module's own shadow reduces the effective ground area from which reflected radiation can "
  "reach the rear surface. The computation therefore requires knowledge of the solar profile "
  "angle, which governs shadow geometry, followed by explicit evaluation of the shadow view "
  "factor F\u1d65.")
P("The solar profile angle \u03b3\u209a\u1d63\u2092\u1da0\u1d62\u2113\u1d49 captures how steeply "
  "the sun appears to climb in the vertical plane perpendicular to the module's long axis. "
  "It is computed from the solar elevation angle and the relative azimuth between sun and "
  "panel [16], [17]:")

def eq6(oM):
    _sub(oM, '\u03b3', 'profile'); _mr(oM, '  =  tan\u207b\u00b9')
    _mr(oM, '\u2061('); 
    _frac(oM, 'tan(sun\u1d49\u2113\u1d49\u1d65)', 'cos(sun\u1d43\u1d63 \u2212 panel\u1d43\u1d63)')
    _mr(oM, ')')
add_eq(eq6, 6)

P("sun\u1d49\u2113\u1d49\u1d65 is the solar elevation angle, sun\u1d43\u1d63 is the solar azimuth "
  "angle, and panel\u1d43\u1d63 is the panel azimuth angle (180\u00b0 for south-facing "
  "installations in the northern hemisphere). The profile angle effectively collapses the "
  "three-dimensional sun position into a single angle relevant to the transverse cross-section "
  "of the module row.")

H("3.2.3  Shadow and View Factor Computation", 12)
P("Using the profile angle, the lower and upper edges of the shadow cast by the module onto "
  "the ground behind it can be located geometrically [16], [17]. If the module lower edge is "
  "at height h above the ground and the module has width W, then:")

def eq7(oM):
    _mr(oM, 'shadow'); _sub(oM, '', 'lower'); _mr(oM, '  =  ')
    _frac(oM, 'h', 'tan(\u03b3\u209a\u1d63\u2092\u1da0\u1d62\u2113\u1d49)')
add_eq(eq7, 7)

def eq8(oM):
    _mr(oM, 'shadow'); _sub(oM, '', 'upper'); _mr(oM, '  =  ')
    _frac(oM, 'h + W sin\u03b2', 'tan(\u03b3\u209a\u1d63\u2092\u1da0\u1d62\u2113\u1d49)')
add_eq(eq8, 8)

P("h (m) is the mounting height of the module lower edge, W (m) is the module width along the "
  "tilt direction, and \u03b2 is the tilt angle. The ground strip between shadow\u2097\u2092\u1d64\u1d49\u1d63 "
  "and shadow\u1d64\u209a\u209a\u1d49\u1d63 receives no direct beam irradiance; only diffuse "
  "sky radiation is reflected from it. The rear sky view factor (rear\u1d65\u1da0) is "
  "geometrically identical to equation (3) with \u03b2 replaced by (180\u00b0 \u2212 \u03b2).")
P("The shadow view factor F\u1d65, introduced by Yusufoglu et al. [18], quantifies the "
  "reduction in rear irradiance attributable to the module shadow. It is expressed as a "
  "surface-integral over the shadow region [18]:")

def eq9(oM):
    _sub(oM, 'F', 'V'); _mr(oM, '  =  rear'); _sub(oM, '', 'vf')
    _mr(oM, '  \u2212  \u222b')
    _frac(oM, 'cos\u03b8\u2081 \u00b7 cos\u03b8\u2082', '2r\u209b\u02b0\u1d43\u1d48\u2092\u1d64')
    _mr(oM, ' dx')
add_eq(eq9, 9)

P("\u03b8\u2081 and \u03b8\u2082 are the angles between the respective surface normals (ground "
  "and module rear face) and the line connecting the differential area elements; r\u209b\u02b0\u1d43\u1d48\u2092\u1d64 "
  "is the distance between those elements, as shown in Fig. 3.2. Incorporating F\u1d65 into "
  "the rear irradiance balance separates the ground-reflected radiation into a diffuse component "
  "(originating from the entire underlying surface) and a direct component (originating only from "
  "the unshaded strip). The net rear irradiance before bifaciality scaling is [18]:")

def eq10(oM):
    _sub(oM, 'G', 'r0'); _mr(oM, '  =  \u03c1 \u00b7 DHI \u00b7 rear')
    _sub(oM, '', 'vf'); _mr(oM, '  +  \u03c1 \u00b7 (GHI \u2212 DHI) \u00b7 (rear')
    _sub(oM, '', 'vf'); _mr(oM, '  \u2212  '); _sub(oM, 'F', 'V'); _mr(oM, ')')
add_eq(eq10, 10)

P("The first term accounts for diffuse-sky-reflected radiation reaching the full rear surface "
  "area; the second term accounts for direct-beam-reflected radiation that can reach only the "
  "unshaded portion. Finally, the effective rear irradiance is scaled by the module's "
  "bifaciality factor \u03c6 to convert rear-surface irradiance into an equivalent front-surface "
  "value for electrical modelling [18]:")

def eq11(oM):
    _sub(oM, 'G', 'rear'); _mr(oM, '  =  '); _sub(oM, 'G', 'r0')
    _mr(oM, '  \u00d7  \u03c6')
add_eq(eq11, 11)

P("The bifaciality factor \u03c6 is the ratio of rear-side to front-side conversion efficiency "
  "measured under standard test conditions (IEC TS 60904-1-2). Commercial bifacial modules "
  "typically exhibit \u03c6 in the range 0.70 to 0.90. A value of \u03c6 = 0.70 is adopted "
  "throughout this work, consistent with conservative field practice for monocrystalline "
  "silicon bifacial modules.")

doc.save('ch3_partA.docx')
print("Part A saved.")

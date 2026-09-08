from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH as A
doc=Document('_part3.docx')
for s in doc.sections: s.top_margin=Inches(1);s.bottom_margin=Inches(1);s.left_margin=Inches(1.5);s.right_margin=Inches(1)
def H(t,sz=14,al=A.LEFT):
    p=doc.add_paragraph();p.alignment=al;p.paragraph_format.space_before=Pt(12);p.paragraph_format.space_after=Pt(6);r=p.add_run(t);r.bold=True;r.font.size=Pt(sz);r.font.name='Times New Roman'
def P(t,al=A.JUSTIFY):
    p=doc.add_paragraph();p.alignment=al;p.paragraph_format.space_after=Pt(6);p.paragraph_format.space_before=Pt(3);r=p.add_run(t);r.font.size=Pt(12);r.font.name='Times New Roman'
def NL():
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(0);p.paragraph_format.space_after=Pt(0)
def FIG(n,c):
    NL();P(f'[Insert {n} here]',al=A.CENTER);P(f'{n}: {c}',al=A.CENTER);NL()
def TBL(headers,rows,caption=None):
    if caption: H(caption,12,A.CENTER)
    t=doc.add_table(rows=len(rows)+1,cols=len(headers));t.style='Table Grid'
    for j,h in enumerate(headers):
        c=t.cell(0,j);c.text='';cp=c.paragraphs[0];cp.alignment=A.CENTER;r=cp.add_run(h);r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(10)
    for i,row in enumerate(rows):
        for j,v in enumerate(row):
            c=t.cell(i+1,j);c.text='';cp=c.paragraphs[0];cp.alignment=A.CENTER;r=cp.add_run(str(v));r.font.name='Times New Roman';r.font.size=Pt(10)

H("CHAPTER 4",16,A.CENTER);H("RESULTS AND DISCUSSION",14,A.CENTER);NL()

H("4.1  Case Study Configuration")
P("To evaluate the proposed framework, a case study was conducted for the location Greater Noida, India (latitude 28.47\u00b0N, longitude 77.50\u00b0E), using the developed frontend-backend analysis platform. Greater Noida was selected as the study location due to its proximity to the project institution and its representative solar resource characteristics for the northern Indian plains. The simulation date was fixed as 13 March 2026 in order to obtain location-specific solar irradiance data for a day near the spring equinox, when the sun path provides a balanced representation of solar geometry conditions.")
P("A parametric sweep was performed over a range of geometric configurations as described below. The module mounting height was varied from 50 cm to 450 cm with increments of 50 cm, providing nine distinct height levels for evaluation. The tilt angle was varied from 10 degrees to 50 degrees in steps of 10 degrees, yielding five tilt angle levels. The azimuth angle of the PV module was fixed at 180 degrees for south-facing orientation, which is the standard configuration for photovoltaic installations in the northern hemisphere to maximize annual solar exposure.")
P("The bifaciality factor was assumed to be 0.7, representing a practical value for commercially available modern bifacial modules. The ground surface type was selected as concrete, corresponding to an albedo value of 0.3, which represents a commonly encountered surface type in urban and peri-urban solar installations. The Nominal Operating Cell Temperature (NOCT) was maintained at 45 degrees Celsius.")
P("The irradiance data was fetched using the NASA POWER API, and the observed total GHI for the selected date was 6.08 kWh/m\u00b2 with a peak GHI of 866.53 W/m\u00b2. These values are consistent with the expected solar resource availability for the Greater Noida region during the month of March, confirming the reliability of the data retrieval process. The DHI component was also retrieved, enabling the decomposition of global irradiance into its beam and diffuse constituents as required by the mathematical formulation described in Chapter 3.")

H("4.2  Combined Parametric Sweep Results")
P("The parametric analyses were executed using the computational engine incorporating the mathematical formulation discussed in Chapter 3. A total of 45 evaluation points (9 height levels \u00d7 5 tilt angle levels) were computed for the concrete surface (albedo = 0.3). For each evaluation point, the framework calculated the total energy output over the analysis day, peak power, front-side irradiance contribution, rear-side irradiance contribution, and the rear gain percentage.")
P("The results revealed clear trends in the variation of system performance with installation parameters. Increasing the mounting height consistently improved the rear-side irradiance contribution across all tilt angles, as higher elevations allow the rear surface to view a larger unshaded ground area. Similarly, the choice of tilt angle affected both the front-side and rear-side performance, with intermediate tilt angles providing the best balance between direct beam capture on the front surface and reflected radiation capture on the rear surface.")
P("The optimal configuration among the 45 evaluation points was identified based on the maximum energy optimization objective. The results indicated an optimal tilt angle of 30 degrees and a mounting height of 450 cm for the selected albedo value of 0.3 (concrete surface). At this optimal configuration, the rear gain was computed to be 16.55% using the formulation described in equations (10) and (11). This result demonstrates that a properly configured bifacial PV installation at the selected location can achieve substantial additional energy generation from the rear surface compared to an equivalent monofacial installation.")

H("4.3  I-V and P-V Characteristics")
P("The current-voltage (I-V) and power-voltage (P-V) characteristics were plotted for the various parameter-based configurations analyzed in the parametric sweep. The I-V curves illustrate the variation in short-circuit current and open-circuit voltage across different installation configurations, reflecting the changes in effective irradiance received by the bifacial module under each scenario.")
P("The short-circuit current showed a strong positive correlation with the total effective irradiance, as expected from the linear relationship between photocurrent generation and incident radiation intensity. Configurations with higher mounting heights and moderate tilt angles exhibited higher short-circuit currents due to the increased rear-side irradiance contribution. The open-circuit voltage showed smaller variations across configurations, consistent with the logarithmic dependence of voltage on irradiance.")
P("The P-V curves clearly show the shift in maximum power point (MPP) as the installation parameters are varied. Configurations with higher effective irradiance (resulting from optimal tilt, height, and albedo combinations) exhibit higher peak power values on the P-V curve. The graphical representation of these characteristics provides a visual confirmation of the parametric analysis results and enables the system designer to assess the electrical performance implications of different installation choices.")
FIG("Figure 4.1","I-V characteristics for different configurations")
FIG("Figure 4.2","P-V characteristics for different configurations")

H("4.4  Optimal Configuration: Maximum Energy")
P("Table 4.1 presents the top five configurations ranked by maximum energy output. The optimization objective in this case is to maximize the total energy generated by the bifacial PV system over the analysis period.")
TBL(["Rank","Height (cm)","Tilt (\u00b0)","Surface","Energy (kWh)","Rear Gain (%)"],
    [["1","450","30","Concrete","3.16843","16.5518"],["2","450","40","Concrete","3.16570","15.6057"],["3","400","30","Concrete","3.16521","16.3917"],["4","400","40","Concrete","3.16258","15.4506"],["5","350","30","Concrete","3.16113","16.1891"]],
    "Table 4.1: Maximum Energy Based Optimal Configuration (Top 5)")
NL()
P("The results show that the highest energy output of 3.168 kWh is achieved at a height of 450 cm and tilt angle of 30 degrees with the concrete surface. The top configurations consistently favor heights of 350\u2013450 cm and tilt angles of 30\u201340 degrees. The rear gain percentages for the top energy configurations range from 15.45% to 16.55%, indicating significant contribution from the rear surface. The optimal tilt angle of 30 degrees closely approximates the latitude of Greater Noida (28.47\u00b0N), which is consistent with the established principle that the optimal tilt angle for annual energy collection is approximately equal to the site latitude.")
P("It is noteworthy that the difference in total energy between the first-ranked and fifth-ranked configurations is only approximately 0.23%, indicating that the system performance is relatively robust to small variations in the installation parameters near the optimal configuration. This finding is practically significant as it implies that minor deviations from the theoretically optimal configuration due to site-specific constraints will have minimal impact on overall system performance.")

H("4.5  Optimal Configuration: Maximum Rear Gain")
P("Table 4.2 presents the top five configurations ranked by maximum rear gain percentage. This optimization objective prioritizes maximizing the relative contribution of the rear surface to the total energy output, which may be relevant for applications where the rear-side performance enhancement is of primary interest.")
TBL(["Rank","Height (cm)","Tilt (\u00b0)","Surface","Energy (kWh)","Rear Gain (%)"],
    [["1","450","10","Concrete","2.99422","17.8989"],["2","400","10","Concrete","2.99130","17.7561"],["3","350","10","Concrete","2.98755","17.5731"],["4","300","10","Concrete","2.98258","17.3301"],["5","450","20","Concrete","3.11109","17.3252"]],
    "Table 4.2: Maximum Rear Gain Based Optimal Configuration (Top 5)")
NL()
P("When the optimization objective is shifted to maximum rear gain, a distinctly different set of optimal configurations emerges. The highest rear gain of 17.90% is achieved at a height of 450 cm and a lower tilt angle of 10 degrees. This is because lower tilt angles increase the ground view factor for the rear surface, allowing more ground-reflected radiation to reach the rear side of the module. The ground view factor for the rear surface increases as the tilt angle decreases because the rear surface becomes more nearly parallel to the ground, thereby viewing a larger ground area.")
P("However, the total energy output for these configurations (2.994 kWh) is approximately 5.5% lower compared to the maximum energy configurations (3.168 kWh), illustrating the trade-off between total energy maximization and rear gain maximization. This trade-off arises because low tilt angles reduce the direct beam component on the front surface, which more than offsets the increased rear-side contribution. The PV system designer can choose the appropriate optimization objective based on their specific requirements and priorities.")

H("4.6  Individual Parameter Variation Analysis")
P("In addition to the combined parametric sweep, the framework provides a provision for fixing two parameters and varying the third to obtain focused analyses of individual parameter effects. This feature makes the platform dynamic and adaptable to the practical needs of PV system designers who may wish to understand the sensitivity of system performance to specific parameters.")
FIG("Figure 4.3","Frontend interface for fixing two parameters and varying third")

H("4.6.1  Effect of Albedo Variation",12)
P("Figure 4.4 shows the variation of rear share of effective irradiance when varying the albedo (surface type) with the following fixed parameters: tilt angle of 20 degrees, height of 100 cm, azimuthal angle of 180 degrees, and bifaciality of 0.7. The results demonstrate that the rear gain share increases monotonically with increasing albedo value, confirming the expected linear relationship between ground reflectivity and rear-side irradiance contribution.")
P("The highest rear gain share of 24.72% was achieved for the aluminum surface (albedo = 0.85), confirming the strong dependence of rear-side performance on ground surface reflectivity. Low-albedo surfaces such as dry asphalt (albedo approximately 0.12) produced rear gain values below 5%, while moderate-albedo surfaces such as concrete (albedo approximately 0.30) yielded intermediate values of approximately 10\u201312%. These findings are consistent with the results reported by Dincer and Ozer [11], who found that aluminum surfaces provided the highest rear-surface energy generation of 21.2%, and by Ganesan et al. [20] who reported an average bifacial gain of 21.4% from aluminum surfaces.")
P("The near-linear relationship between albedo and rear gain observed in the results is consistent with the mathematical formulation in equation (10), where the ground-reflected irradiance is directly proportional to the albedo coefficient. This linear dependency has been independently confirmed by Yusufoglu et al. [18] in their simulation study and validates the implementation of the rear-side irradiance model in the proposed framework.")
FIG("Figure 4.4","Rear share of effective irradiance on varying albedo")

H("4.6.2  Effect of Height Variation",12)
P("Figure 4.5 illustrates the variation of rear share of effective irradiance when varying the panel height with the tilt angle fixed at 20 degrees, albedo of 0.18, and all other factors remaining constant. The results show that rear gain increases with panel height, achieving 6.61% for a panel height of 100 cm. The rate of increase diminishes at higher elevations, consistent with the near-logarithmic dependence of rear-side performance on module elevation demonstrated by Yusufoglu et al. [18].")
P("This behavior is expected because the geometric view factor between the module and the ground improves with height but approaches an asymptotic limit as the module elevation increases beyond a certain threshold relative to the module dimensions. At very high elevations, nearly the entire ground surface is visible to the rear of the module, and further height increases provide negligible improvement in the view factor. The practical implication of this finding is that there exists an optimal range of mounting heights beyond which the additional structural costs and complexity of higher mounting systems cannot be justified by the marginal energy gains.")
P("The height dependence observed in this study is quantitatively consistent with the results reported by Dincer and Ozer [11], who found that the rear gain increased from 4.1% at 40 cm to 4.5% at 100 cm for urban albedo conditions. The slightly higher values obtained in the present study can be attributed to differences in the irradiance conditions and the specific albedo value used.")
FIG("Figure 4.5","Rear share of effective irradiance on varying height")

H("4.6.3  Effect of Tilt Angle Variation",12)
P("Figure 4.6 presents the variation of rear share of effective irradiance when varying the tilt angle with the height fixed at 100 cm, albedo of 0.18, and all other factors remaining constant. The results indicate a maximum rear gain share of 7.24% at a tilt angle of 10 degrees. The rear gain decreases with increasing tilt angle in this configuration.")
P("This trend can be attributed to the change in balance between sky view and ground view factors for the rear surface. At lower tilt angles, the ground view factor for the rear surface is larger, allowing more ground-reflected radiation to reach the rear side. As the tilt angle increases, the rear ground view factor decreases while the rear sky view factor increases. Since the ground-reflected radiation (which depends on albedo) typically contributes more to rear-side irradiance than the diffuse sky radiation at moderate to high albedo values, lower tilt angles tend to maximize the rear-side contribution.")
P("However, the total system energy output is not maximized at the lowest tilt angle because the front-side direct beam irradiance capture is reduced at low tilt angles for locations at moderate to high latitudes. This finding further reinforces the importance of considering both front-side and rear-side performance simultaneously when optimizing bifacial PV installations, and highlights the value of the dual optimization capability provided by the proposed framework.")
FIG("Figure 4.6","Rear share of effective irradiance on varying tilt angle")

H("4.7  Comparison with Published Literature")
P("To validate the proposed framework, the results obtained in this study were compared with findings reported in published literature. Table 4.3 presents a systematic comparison of key results.")
TBL(["Study","Scenario","Reported Value","This Study"],
    [["Dincer & Ozer [11]","Aluminum surface, tilt 20\u00b0, h=100cm","21.2% rear gain","24.72% rear gain"],
     ["Dincer & Ozer [11]","Height 40-100cm, urban albedo","4.1-4.5% rear gain","5.2-6.6% rear gain"],
     ["Dincer & Ozer [11]","Tilt 10\u00b0-50\u00b0 variation","4.3-5.5% rear gain","4.1-7.2% rear gain"],
     ["Yusufoglu et al. [18]","Albedo=0.2, 2m height, Cairo","~13.8% energy gain","16.55% rear gain"],
     ["Ganesan et al. [20]","Aluminum, bifacial gain","21.4% avg gain","24.72% rear gain"],
     ["Pelaez et al. [21]","High albedo validation","~20% bifacial gain","17.9-24.7% range"]],
    "Table 4.3: Comparison of Results with Published Literature")
NL()
P("The comparison reveals that the results obtained using the proposed framework are consistent with the trends and magnitudes reported in the reference studies. The slightly higher values obtained in some cases can be attributed to differences in geographic location (and hence solar resource characteristics), assumed panel parameters, temporal resolution of irradiance data, and specific modeling methodology. The overall agreement validates the reliability and accuracy of the proposed computational framework for bifacial PV system optimization.")
P("In particular, the albedo sensitivity results show excellent qualitative agreement with the literature, with the ranking of surface types by rear gain being consistent across all studies. The height dependency trend (logarithmic increase with diminishing returns) is also consistent with the simulation results of Yusufoglu et al. [18]. The tilt angle sensitivity results align with the findings of Dincer and Ozer [11], confirming that rear gain increases at lower tilt angles due to enhanced ground view factors.")

H("4.8  Discussion of Results")
P("The parametric analysis results provide several important insights for the design and optimization of bifacial PV installations:")
P("First, the analysis confirms that ground surface albedo is the most influential parameter affecting rear-side irradiance contribution. The rear gain varied from less than 5% for low-albedo surfaces to over 24% for highly reflective surfaces, representing a variation of nearly five-fold. This finding has direct practical implications, suggesting that site preparation measures such as installing reflective ground cover materials (white paint, gravel, or reflective membranes) can substantially improve the performance of bifacial PV systems at relatively low additional cost.")
P("Second, the results demonstrate that mounting height has a positive but diminishing effect on rear-side performance. While higher installations consistently yield higher rear gains, the incremental improvement decreases with increasing elevation. This suggests that there exists a practical optimum height beyond which the additional structural costs and complexity of higher mounting systems cannot be justified by the marginal energy gains. For the case study conditions, mounting heights in the range of 200\u2013300 cm appear to provide a good balance between performance improvement and practical feasibility.")
P("Third, the analysis reveals that the optimal tilt angle depends on the chosen optimization objective. When maximizing total energy, the optimal tilt angle (30\u00b0 for Greater Noida) closely corresponds to the latitude, similar to monofacial modules. When maximizing rear gain, lower tilt angles (10\u201320\u00b0) are preferred due to enhanced ground view factors. This dual optimization capability of the proposed framework enables system designers to make informed decisions based on their specific priorities.")
P("Fourth, the combined parametric sweep approach provides more comprehensive optimization compared to single-parameter studies. The interaction effects between tilt angle, height, and albedo are captured in the multi-dimensional analysis, leading to more accurate identification of globally optimal configurations. The relatively small performance variation among the top-ranked configurations (less than 0.25% difference in energy) indicates that the system performance is robust to minor installation variations near the optimum.")
P("Fifth, the comparison with published literature confirms the validity of the proposed framework and demonstrates that the mathematical formulation and computational implementation produce results consistent with established research in the field.")

doc.add_page_break()
doc.save('_part4.docx')
print("Ch4 done.")

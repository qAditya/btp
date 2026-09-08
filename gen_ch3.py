from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document('BTP_Report_Ch1Ch2.docx')
for s in doc.sections: s.top_margin=Inches(1); s.bottom_margin=Inches(1); s.left_margin=Inches(1.5); s.right_margin=Inches(1)
def H(t,sz=14,al=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_paragraph(); p.alignment=al; r=p.add_run(t); r.bold=True; r.font.size=Pt(sz); r.font.name='Times New Roman'
def P(t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.space_after=Pt(6); r=p.add_run(t); r.font.size=Pt(12); r.font.name='Times New Roman'
def NL(): p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)

H("CHAPTER 3",16,WD_ALIGN_PARAGRAPH.CENTER)
H("METHODOLOGY",14,WD_ALIGN_PARAGRAPH.CENTER)
NL()
H("3.1  Overview of the Proposed Framework")
P("This chapter presents the detailed methodology adopted for the parametric analysis and optimization of bifacial photovoltaic systems. The proposed framework integrates three core components: (i) a mathematical formulation for computing the front-side and rear-side irradiance on bifacial PV modules, (ii) a MATLAB/Simulink model for simulating the electrical characteristics of the PV system, and (iii) a user-oriented frontend platform that retrieves location-specific irradiance data and performs automated parametric sweep analyses. The framework is designed to evaluate and determine the optimal configuration of bifacial PV systems by varying key installation parameters including module tilt angle, ground surface albedo, and mounting height. The computational pipeline begins with the user specifying the geographical location, date range, and parameter sweep ranges through the frontend interface. The backend application processes these inputs, retrieves hourly irradiance data from the NASA POWER API, and executes the analysis engine to compute the system performance metrics for each combination of installation parameters. The results are presented through an interactive dashboard displaying the optimal configurations, I-V and P-V characteristics, and comparative performance metrics.")

H("3.2  Mathematical Formulation of Bifacial PV Operation")
P("Unlike conventional monofacial PV modules, which receive irradiance only on the front surface, bifacial modules are capable of utilizing additional radiation reflected from the ground and surrounding surfaces. The total effective irradiance on a bifacial module can therefore be expressed as the sum of the front and rear irradiance components [4]:")
P("G_tot = G_front + G_rear    ... (1)")
P("where G_tot is the total irradiance received by the bifacial module, G_front is the irradiance incident on the front surface, and G_rear is the irradiance incident on the rear surface. The front and rear irradiance components depend on several factors which are taken into account in the analysis as described in the following subsections.")

H("3.2.1  Front-Side Irradiance Model",12)
P("The front-side irradiance is estimated using the Liu-Jordan isotropic sky model [15], which decomposes the global horizontal irradiance into beam, diffuse, and ground-reflected components. The direct beam irradiance incident on the tilted module surface is calculated as:")
P("beam_tilted = DNI x max(0, cos I)    ... (2)")
P("where DNI represents the direct normal irradiance and I denotes the angle of incidence between the incoming solar radiation and the module surface. The max(0, cos I) term ensures that negative irradiance values are avoided when the sun is below the panel horizon.")
P("The diffuse irradiance component is modeled assuming isotropic sky distribution. The sky view factor and ground view factor for the tilted surface are given by [15]:")
P("sky_vf = (1 + cos beta) / 2    ... (3)")
P("gnd_vf = (1 - cos beta) / 2    ... (4)")
P("where beta represents the module tilt angle. The sky view factor quantifies the portion of the sky dome visible to the tilted PV surface, while the ground view factor represents the fraction of ground-reflected radiation reaching the module surface. Combining these contributions, the total front-side irradiance is calculated as [15]:")
P("G_front = beam_tilted + (DHI x sky_vf) + (GHI x rho x gnd_vf)    ... (5)")
P("where DHI is the diffuse horizontal irradiance, GHI is the global horizontal irradiance, and rho is the ground surface albedo coefficient.")

H("3.2.2  Rear-Side Irradiance Model",12)
P("The computation of rear irradiance depends on the shadow view factor F_V, albedo, DHI, GHI, and rear sky view factor. The shadow view factor F_V represents the reduction in rear irradiance due to row-to-row shading in the PV array. It is determined by evaluating the solar profile angle and the geometric relationship between the sun position and the module orientation.")
P("The solar profile angle is computed as [16], [17]:")
P("gamma_profile = arctan(tan(sun_elev) / cos(sun_az - panel_az))    ... (6)")
P("where sun_elev is the solar elevation angle, and sun_az and panel_az represent the solar and panel azimuth angles, respectively.")

H("3.2.3  Shadow and View Factor Computation",12)
P("The lower and upper bounds of the shadow region are calculated using the solar profile angle as follows [16], [17]:")
P("shadow_lower = h / tan(gamma_profile)    ... (7)")
P("shadow_upper = (h + W sin beta) / tan(gamma_profile)    ... (8)")
P("where h is the mounting height and W is the width of the module. The rear sky view factor (rear_vf) is geometrically represented by equation (3) with beta replaced by (180 - beta). The shadow view factor F_V is calculated by integrating over the shadow region as described by Yusufoglu et al. [18]:")
P("F_V = rear_vf - integral(cos(theta_1) cos(theta_2) / (2r_shadow)) dx    ... (9)")
P("where theta_1 and theta_2 represent the geometric angles between the ground and module surfaces, and r is the distance between the differential elements. The final rear irradiance is given by incorporating the above factors [18]:")
P("G_r0 = rho . DHI . rear_vf + rho . (GHI - DHI) . (rear_vf - F_V)    ... (10)")
P("This formulation separates the ground-reflected radiation into diffuse and direct components. The diffuse component (DHI) is reflected from the entire ground surface, while the direct component (GHI - DHI) is reflected only from the unshadowed portion of the ground. The effective rear irradiance is then scaled by the bifaciality factor:")
P("G_rear = G_r0 x bifaciality    ... (11)")
P("The bifaciality factor represents the ratio of rear-side efficiency to front-side efficiency of the bifacial PV module and is typically in the range of 0.70 to 0.90 for modern bifacial modules. In this study, a bifaciality factor of 0.7 is assumed based on practical considerations.")

H("3.3  Simulink Model Development")
P("The bifacial photovoltaic system is modeled in MATLAB/Simulink to analyze the electrical performance of the module under varying environmental and installation conditions. The model incorporates key parameters such as global horizontal irradiance (GHI), ground reflectivity (albedo), tilt angle, mounting height, and temperature.")
P("The input parameters are processed through a MATLAB function block (calculate_irradiance), which determines the effective irradiance G_eff received by the module using equation (1). The resulting effective irradiance value is then supplied to the PV Array block as the irradiance input. The PV Array block represents the electrical model of the photovoltaic module and produces the output current and voltage corresponding to the applied operating conditions.")
P("To obtain the electrical characteristics of the PV system, a controlled voltage source is used to perform an I-V sweep across the PV module terminals. The voltage sweep is generated using a ramp signal (V_ramp), which gradually varies the applied voltage across the PV array over the simulation interval. As the voltage changes, the PV module produces corresponding current values, allowing the complete I-V characteristic curve of the module to be obtained. The measured current (I_meas) and voltage (V_meas) signals are collected through a bus creator block and forwarded to subsequent processing stages. The first scope plots the current-voltage (I-V) characteristics, while the second scope plots the power-voltage (P-V) characteristics of the bifacial PV module.")
P("[Figure 3.3: Simulink model for Bifacial PV Module - Insert Simulink screenshot here]")

H("3.4  Frontend Platform Architecture")
P("The proposed framework is designed as a multi-layered system consisting of four primary components: the user interface, backend application, analysis and simulation engine, and the results dashboard.")

H("3.4.1  User Interface Layer",12)
P("The process begins at the user interface, where the user provides the necessary input parameters required for the analysis. These inputs include the geographical location in the form of a selected city, the desired date range for analysis, and key system configuration parameters such as module tilt angle range, mounting height range, and ground surface type. The interface is designed to be intuitive and accessible, requiring no specialized knowledge of PV system modeling or solar engineering.")

H("3.4.2  Backend Application Layer",12)
P("Once the input parameters are specified, the backend application processes these inputs and generates a request for solar irradiance data. The irradiance data is obtained through the NASA POWER API, which provides location-specific solar radiation data for the selected city and time period. The backend handles API communication, data parsing, error handling, and the coordination of parametric sweep computations.")

H("3.4.3  Analysis and Simulation Engine",12)
P("The irradiance dataset is passed to the analysis and simulation engine, which performs the computational evaluation of the bifacial PV system. This module integrates the irradiance data with the user-defined installation parameters to simulate the electrical behavior of the bifacial PV system. The simulation process evaluates the I-V and P-V characteristics and calculates key performance indicators such as output power and bifacial gain under different configuration settings. The engine implements the mathematical formulation described in Section 3.2 and performs the parametric sweep across all specified combinations of tilt angle, albedo, and mounting height.")

H("3.4.4  Results Dashboard",12)
P("The computed results are presented through a results dashboard, which provides a structured visualization of the system performance. The dashboard displays the maximum power output and bifacial gain values in tabular form, along with graphical representations of the I-V and P-V characteristics. The dashboard also identifies and highlights the optimal configuration based on the user's selected optimization objective (maximum energy or maximum rear gain).")

H("3.5  NASA POWER API Integration")
P("The NASA Prediction of Worldwide Energy Resources (POWER) project provides access to meteorological and solar radiation data derived from satellite observations and atmospheric models. The POWER API serves as the primary data source for location-specific irradiance information in the proposed framework. For a given geographic location specified by latitude and longitude coordinates, the API returns hourly values of GHI, DHI, DNI, temperature, and other relevant meteorological parameters. The API is accessed through HTTP requests, and the returned JSON data is parsed and processed by the backend application to extract the required irradiance components for the bifacial PV performance analysis.")

H("3.6  Parametric Sweep Configuration")
P("The parametric sweep is configured to systematically evaluate a grid of installation parameter combinations. In the present implementation, the following parameter ranges and increments are used for the case study: module mounting height is varied from 50 cm to 450 cm with increments of 50 cm (9 levels); tilt angle is varied from 10 degrees to 50 degrees in steps of 10 degrees (5 levels); and surface albedo is selected from predefined surface types with known albedo values. This configuration results in 45 evaluation points (9 heights x 5 tilt angles) for each selected albedo value. For each evaluation point, the analysis engine computes the total energy output, peak power, and rear gain percentage. The results are ranked according to the selected optimization objective, and the top configurations are presented to the user. Additionally, the framework provides a provision for fixing two parameters and varying the third to obtain a focused analysis of individual parameter effects, making the platform dynamic and adaptable to the practical needs of PV system designers.")

P("[Figure 3.4: Proposed Framework for Optimal Configuration of Bifacial PV Systems - Insert framework diagram here]")
doc.add_page_break()
doc.save('BTP_Report_Ch3.docx')
print("Chapter 3 saved!")

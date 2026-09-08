"""Build full expanded report with proper equations. Run: python r_build.py"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin=Inches(1);s.bottom_margin=Inches(1);s.left_margin=Inches(1.5);s.right_margin=Inches(1)
st=doc.styles['Normal'];st.font.name='Times New Roman';st.font.size=Pt(12);st.paragraph_format.line_spacing=1.5

def H(t,sz=14,al=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_paragraph();p.alignment=al;p.paragraph_format.space_before=Pt(12);p.paragraph_format.space_after=Pt(6)
    r=p.add_run(t);r.bold=True;r.font.size=Pt(sz);r.font.name='Times New Roman';return p
def P(t,b=False,it=False,al=WD_ALIGN_PARAGRAPH.JUSTIFY,sz=12):
    p=doc.add_paragraph();p.alignment=al;p.paragraph_format.space_after=Pt(6);p.paragraph_format.space_before=Pt(3)
    r=p.add_run(t);r.font.size=Pt(sz);r.font.name='Times New Roman';r.bold=b;r.italic=it;return p
def NL(n=1):
    for _ in range(n):
        p=doc.add_paragraph();p.paragraph_format.space_before=Pt(0);p.paragraph_format.space_after=Pt(0)
def EQ(text, num):
    """Proper equation: Cambria Math font, centered, with number."""
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(8);p.paragraph_format.space_after=Pt(8)
    r=p.add_run(text);r.font.name='Cambria Math';r.font.size=Pt(13);r.italic=True
    r2=p.add_run(f'          ({num})');r2.font.name='Times New Roman';r2.font.size=Pt(12)
def TBL(headers,rows,caption=None):
    if caption: H(caption,12,WD_ALIGN_PARAGRAPH.CENTER)
    t=doc.add_table(rows=len(rows)+1,cols=len(headers));t.style='Table Grid'
    for j,h in enumerate(headers):
        c=t.cell(0,j);c.text='';cp=c.paragraphs[0];cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=cp.add_run(h);r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(10)
    for i,row in enumerate(rows):
        for j,v in enumerate(row):
            c=t.cell(i+1,j);c.text='';cp=c.paragraphs[0];cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=cp.add_run(str(v));r.font.name='Times New Roman';r.font.size=Pt(10)
def FIG(num,cap):
    NL();P(f'[Insert {num} here]',al=WD_ALIGN_PARAGRAPH.CENTER,it=True)
    P(f'{num}: {cap}',b=True,al=WD_ALIGN_PARAGRAPH.CENTER,sz=11);NL()

# ===================== TITLE PAGE =====================
NL(3)
H('B.Tech. Project',16,WD_ALIGN_PARAGRAPH.CENTER)
H('On',14,WD_ALIGN_PARAGRAPH.CENTER)
NL()
H('OPTIMAL CONFIGURATION OF BIFACIAL\nPHOTOVOLTAIC MODULES USING\nPARAMETRIC ANALYSIS',16,WD_ALIGN_PARAGRAPH.CENTER)
NL()
P('Report submitted in partial fulfillment of the requirements for the\nB. Tech. degree in Electrical Engineering',al=WD_ALIGN_PARAGRAPH.CENTER)
NL()
H('By',14,WD_ALIGN_PARAGRAPH.CENTER)
studs=[('Ashish Upadhyay','2022UEE4521'),('Pratyai Chakrabarty','2022UEE4586'),('Aman','2022UEE4532'),('Aditya Garg','2022UEE4503')]
t=doc.add_table(rows=5,cols=2);t.alignment=WD_ALIGN_PARAGRAPH.CENTER
for j,h in enumerate(['Name of Student','Roll No.']):
    c=t.cell(0,j);c.text='';cp=c.paragraphs[0];cp.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=cp.add_run(h);r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12)
for i,(n,ro) in enumerate(studs):
    for j,v in enumerate([n,ro]):
        c=t.cell(i+1,j);c.text='';cp=c.paragraphs[0];cp.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=cp.add_run(v);r.font.name='Times New Roman';r.font.size=Pt(12)
NL()
P('Under the supervision of',al=WD_ALIGN_PARAGRAPH.CENTER)
H('Dr. Astitva Kumar',14,WD_ALIGN_PARAGRAPH.CENTER)
NL(2)
P('[NSUT Logo]',b=True,al=WD_ALIGN_PARAGRAPH.CENTER)
NL()
H('DEPARTMENT OF ELECTRICAL ENGINEERING',14,WD_ALIGN_PARAGRAPH.CENTER)
H('NETAJI SUBHAS UNIVERSITY OF TECHNOLOGY (NSUT)\nDWARKA, NEW DELHI',14,WD_ALIGN_PARAGRAPH.CENTER)
H('MAY 2026',14,WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ===================== DECLARATION =====================
H("CANDIDATE(S) DECLARATION",14,WD_ALIGN_PARAGRAPH.CENTER);NL()
H("DEPARTMENT OF ELECTRICAL ENGINEERING",12,WD_ALIGN_PARAGRAPH.CENTER);NL()
P('I/We, Ashish Upadhyay (2022UEE4521), Pratyai Chakrabarty (2022UEE4586), Aman (2022UEE4532) and Aditya Garg (2022UEE4503), students of B. Tech., Department of Electrical Engineering, hereby declare that the Project-Thesis titled "Optimal Configuration of Bifacial Photovoltaic Modules using Parametric Analysis" which is submitted by us to the Department of Electrical Engineering, Netaji Subhas University of Technology (NSUT) Dwarka, New Delhi in partial fulfillment of the requirement for the award of the degree of Bachelor of Technology is our original work and not copied from any source without proper citation. The manuscript has been subjected to plagiarism check by Turnitin software. This work has not previously formed the basis for the award of any other Degree.')
NL(2);P('Place: New Delhi');P('Date: May 2026');NL(2)
for n,ro in studs: P(f'{n}\n{ro}')
P('\ni',al=WD_ALIGN_PARAGRAPH.CENTER);doc.add_page_break()

# ===================== CERTIFICATE =====================
H("CERTIFICATE OF DECLARATION",14,WD_ALIGN_PARAGRAPH.CENTER);NL()
H("DEPARTMENT OF ELECTRICAL ENGINEERING",12,WD_ALIGN_PARAGRAPH.CENTER);NL()
P('This is to certify that the work embodied in project thesis titled, "Optimal Configuration of Bifacial Photovoltaic Modules using Parametric Analysis" by Ashish Upadhyay (2022UEE4521), Pratyai Chakrabarty (2022UEE4586), Aman (2022UEE4532) and Aditya Garg (2022UEE4503) is the bonafide work of the group submitted to Netaji Subhas University of Technology for consideration in 8th Semester B.Tech. Project Evaluation.')
NL()
P('The original Research work was carried out by the team under my guidance and supervision in the academic year 2025-2026. This work has not been submitted for any other diploma or degree of any university. On the basis of declaration made by the group, we recommend the project report for evaluation.')
NL(3);P('Dr. Astitva Kumar');P('Assistant Professor');P('Department of Electrical Engineering');P('Netaji Subhas University of Technology')
P('\nii',al=WD_ALIGN_PARAGRAPH.CENTER);doc.add_page_break()

# ===================== ACKNOWLEDGEMENT =====================
H("ACKNOWLEDGEMENT",14,WD_ALIGN_PARAGRAPH.CENTER);NL()
P('We would like to express our sincere gratitude and appreciation to all those who made it possible to complete this project. First and foremost, we extend our deepest thanks to our project supervisor, Dr. Astitva Kumar, Assistant Professor, Department of Electrical Engineering, Netaji Subhas University of Technology, New Delhi, whose continuous guidance, invaluable suggestions, constructive criticism, and constant encouragement have been instrumental throughout the course of this work. His expertise in the domain of photovoltaic systems and renewable energy provided the foundation upon which this research was built.')
P('We would also like to acknowledge the Department of Electrical Engineering, NSUT, for providing access to laboratory facilities, computational resources, and licensed software tools including MATLAB/Simulink, which were essential for the simulation and analysis components of this project.')
P('We are grateful to our colleagues and peers who offered their time for technical discussions, proofreading, and constructive feedback during the preparation of this report. Their support has been invaluable in refining both the technical content and the presentation quality of this work.')
P('Finally, we wish to express our heartfelt gratitude to our families for their unwavering support, patience, and motivation throughout the duration of this project.')
NL(2)
for n,_ in studs: P(n,al=WD_ALIGN_PARAGRAPH.RIGHT)
P('\niii',al=WD_ALIGN_PARAGRAPH.CENTER);doc.add_page_break()

# ===================== PLAGIARISM =====================
H("PLAGIARISM REPORT",14,WD_ALIGN_PARAGRAPH.CENTER);NL(2)
P('[Attach Turnitin/Ouriginal plagiarism report here]',b=True,al=WD_ALIGN_PARAGRAPH.CENTER)
P('Similarity Index: ___ % (must be less than 20%)',al=WD_ALIGN_PARAGRAPH.CENTER)
NL(2);P('Verified by:');NL(2);P('Dr. Astitva Kumar');P('Supervisor')
P('\niv',al=WD_ALIGN_PARAGRAPH.CENTER);doc.add_page_break()

# ===================== ABSTRACT =====================
H("ABSTRACT",14,WD_ALIGN_PARAGRAPH.CENTER);NL()
P('Solar photovoltaic (PV) technology has become a key component in the global transition toward sustainable energy, with countries such as India witnessing rapid growth in solar power deployment due to abundant solar resources and increasing energy demand. Efficient utilization of available solar irradiance is therefore critical to maximize energy generation from installed PV systems. In this context, bifacial photovoltaic modules have gained significant attention as they are capable of capturing solar irradiance from both the front and rear surfaces, leading to higher energy yield compared to conventional monofacial panels. However, the performance advantage of bifacial PV systems strongly depends on installation parameters such as module tilt angle, ground surface albedo, and mounting height, which influence the amount of reflected and incident irradiance received by the rear side of the module. This project presents a computational framework for determining the optimal configuration of bifacial PV installations for a given geographical location. A user-oriented front-end platform is developed where inputs such as city and date are provided, and corresponding hourly irradiance data are retrieved from the NASA POWER database. Parametric analyses are then performed by varying key installation parameters including tilt angle, surface albedo, and mounting height. The framework evaluates the resulting energy output and identifies the configuration that maximizes power generation for the selected location. The proposed approach enables location-specific optimization of bifacial PV systems and provides a practical tool for improving system design and deployment strategies.')
NL()
P('Keywords: Solar Photovoltaic, Bifacial PV, Parametric Analysis, Frontend Platform, Strategic Installation, View Factor',b=True)
P('\nv',al=WD_ALIGN_PARAGRAPH.CENTER);doc.add_page_break()

# ===================== TOC placeholder =====================
H("TABLE OF CONTENTS",14,WD_ALIGN_PARAGRAPH.CENTER);NL()
P('[Auto-generate Table of Contents in Microsoft Word:\nReferences > Table of Contents > Automatic Table]',al=WD_ALIGN_PARAGRAPH.CENTER,it=True)
NL()
P('Candidate(s) Declaration ...................................... i')
P('Certificate of Declaration ...................................... ii')
P('Acknowledgement ...................................... iii')
P('Plagiarism Report ...................................... iv')
P('Abstract ...................................... v')
P('Table of Contents ...................................... vi')
P('List of Figures ...................................... ix')
P('List of Tables ...................................... x')
P('List of Abbreviations ...................................... xi')
NL()
P('CHAPTER 1: INTRODUCTION ...................................... 1',b=True)
P('    1.1 Background and Motivation ...................................... 1')
P('    1.2 Solar Photovoltaic Technology: An Overview ...................................... 3')
P('    1.3 Monofacial versus Bifacial PV Modules ...................................... 5')
P('    1.4 Key Parameters Affecting Bifacial PV Performance ...................................... 6')
P('    1.5 Challenges in Bifacial PV Optimization ...................................... 8')
P('    1.6 Problem Statement and Research Objectives ...................................... 9')
P('    1.7 Organization of the Report ...................................... 10')
NL()
P('CHAPTER 2: LITERATURE REVIEW ...................................... 11',b=True)
P('    2.1 Review of Bifacial PV Technology ...................................... 11')
P('    2.2 Influence of Tilt Angle on Bifacial PV Performance ...................................... 13')
P('    2.3 Effect of Ground Albedo on Rear-Side Irradiance ...................................... 15')
P('    2.4 Impact of Mounting Height on Energy Yield ...................................... 17')
P('    2.5 Combined Parametric Optimization Studies ...................................... 19')
P('    2.6 Simulation and Modeling Approaches ...................................... 20')
P('    2.7 Research Gap Analysis ...................................... 22')
NL()
P('CHAPTER 3: METHODOLOGY ...................................... 24',b=True)
P('    3.1 Overview of the Proposed Framework ...................................... 24')
P('    3.2 Mathematical Formulation of Bifacial PV Operation ...................................... 25')
P('    3.3 Simulink Model Development ...................................... 31')
P('    3.4 Frontend Platform Architecture ...................................... 33')
P('    3.5 NASA POWER API Integration ...................................... 35')
P('    3.6 Parametric Sweep Configuration ...................................... 36')
NL()
P('CHAPTER 4: RESULTS AND DISCUSSION ...................................... 37',b=True)
P('    4.1 Case Study Configuration ...................................... 37')
P('    4.2 Combined Parametric Sweep Results ...................................... 38')
P('    4.3 I-V and P-V Characteristics ...................................... 39')
P('    4.4 Optimal Configuration: Maximum Energy ...................................... 40')
P('    4.5 Optimal Configuration: Maximum Rear Gain ...................................... 41')
P('    4.6 Individual Parameter Variation Analysis ...................................... 42')
P('    4.7 Comparison with Published Literature ...................................... 45')
P('    4.8 Discussion of Results ...................................... 47')
NL()
P('CHAPTER 5: CONCLUSION AND FUTURE WORK ...................................... 49',b=True)
P('    5.1 Summary of Contributions ...................................... 49')
P('    5.2 Key Findings ...................................... 50')
P('    5.3 Limitations ...................................... 51')
P('    5.4 Scope for Future Work ...................................... 52')
NL()
P('REFERENCES ...................................... 53',b=True)
P('APPENDIX A ...................................... 56',b=True)
P('APPENDIX B ...................................... 58',b=True)
P('BIO-DATA ...................................... 60',b=True)
doc.add_page_break()

# ===================== LIST OF FIGURES =====================
H("LIST OF FIGURES",14,WD_ALIGN_PARAGRAPH.CENTER);NL()
figs=[("1.1","Global solar PV installed capacity growth (2015-2025)"),("1.2","Schematic of monofacial vs bifacial PV modules"),("1.3","Key parameters influencing bifacial PV module performance"),("2.1","Year-wise publication trend on bifacial PV research"),("2.2","Albedo values of common ground surfaces"),("3.1","Key parameters affecting bifacial PV modules [11]"),("3.2","View factor representation for bifacial PV"),("3.3","Simulink model for bifacial PV module"),("3.4","Proposed framework for optimal configuration"),("3.5","Frontend user interface screenshot"),("4.1","I-V characteristics for different configurations"),("4.2","P-V characteristics for different configurations"),("4.3","Frontend interface for fixing two parameters"),("4.4","Rear share of effective irradiance on varying albedo"),("4.5","Rear share of effective irradiance on varying height"),("4.6","Rear share of effective irradiance on varying tilt angle")]
for fn,fc in figs: P(f'Figure {fn}: {fc}')
doc.add_page_break()

# ===================== LIST OF TABLES =====================
H("LIST OF TABLES",14,WD_ALIGN_PARAGRAPH.CENTER);NL()
for tn,tc in [("2.1","Albedo values of common surface types"),("2.2","Summary of key bifacial PV parametric studies"),("3.1","Technical specifications of the PV module"),("4.1","Max energy based optimal configuration"),("4.2","Max rear gain based optimal configuration"),("4.3","Comparison with published literature")]:
    P(f'Table {tn}: {tc}')
doc.add_page_break()

# ===================== ABBREVIATIONS =====================
H("LIST OF ABBREVIATIONS",14,WD_ALIGN_PARAGRAPH.CENTER);NL()
for a,f in [("PV","Photovoltaic"),("BPV","Bifacial Photovoltaic"),("GHI","Global Horizontal Irradiance"),("DHI","Diffuse Horizontal Irradiance"),("DNI","Direct Normal Irradiance"),("API","Application Programming Interface"),("NASA","National Aeronautics and Space Administration"),("POWER","Prediction of Worldwide Energy Resources"),("NOCT","Nominal Operating Cell Temperature"),("STC","Standard Test Conditions"),("I-V","Current-Voltage"),("P-V","Power-Voltage"),("MPP","Maximum Power Point"),("PERC","Passivated Emitter and Rear Contact"),("TOPCon","Tunnel Oxide Passivated Contact"),("HJT","Heterojunction Technology"),("IEA","International Energy Agency"),("NREL","National Renewable Energy Laboratory")]:
    p=doc.add_paragraph();r1=p.add_run(a);r1.bold=True;r1.font.name='Times New Roman';r1.font.size=Pt(12)
    r2=p.add_run(f'    {f}');r2.font.name='Times New Roman';r2.font.size=Pt(12)
doc.add_page_break()

doc.save('_part0.docx')
print("Part 0 (front matter) done.")

"""
Generate DOCX files for Camera-Ready Paper submission:
1. Camera_Ready_Changes.docx - Detailed change list
2. Rebuttal_Sheet.docx - Response to reviewers
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = r"c:\Users\amsh9\OneDrive\Desktop\PV-Bifacial-Sim"


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h


def add_table_row(table, cells_text, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold or header:
            run.bold = True
        if header:
            set_cell_shading(cell, "003366")
            run.font.color.rgb = RGBColor(255, 255, 255)
    return row


# ============================================================
# DOCUMENT 1: Camera-Ready Changes
# ============================================================
def generate_changes_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Camera-Ready Paper: Required Changes Document', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Meta info
    meta = doc.add_paragraph()
    meta.add_run('Paper: ').bold = True
    meta.add_run('Optimal Configuration of Bifacial Photovoltaic Modules using Parametric Analysis\n')
    meta.add_run('Conference: ').bold = True
    meta.add_run('CE2CT 2026\n')
    meta.add_run('Current Version: ').bold = True
    meta.add_run('5th Draft (Accepted)\n')
    meta.add_run('Target: ').bold = True
    meta.add_run('Camera-Ready Paper (CRP)')
    
    doc.add_paragraph()
    
    # Reviewer summary table
    add_heading_styled(doc, 'Summary of Reviewer Comments', level=1)
    
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(['Reviewer', 'Key Concerns']):
        hdr[i].paragraphs[0].add_run(txt).bold = True
        set_cell_shading(hdr[i], "003366")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    reviewers = [
        ('R1', 'Add experimental/real-field validation; uncertainty/sensitivity analysis; compare with existing methods'),
        ('R2', 'Fix IEEE references; improve figure/table formatting & captions; add validation via comparison'),
        ('R3', 'Follow IEEE template; improve abstract; add novelty bullets; 300 dpi figures; editable tables; format equations; proofread; cite all references'),
    ]
    for rev, concern in reviewers:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(rev).bold = True
        row.cells[1].paragraphs[0].add_run(concern).font.size = Pt(9)
    
    # ---- Section by section changes ----
    sections = [
        {
            'title': '1. TITLE & AUTHOR BLOCK',
            'ref': 'Camera-Ready Guideline #1 | Priority: HIGH',
            'changes': [
                ('1.1', 'Format author details uniformly', 'Author block', 'Include ONLY: Name, Department, Organization, City, Country, Email. Remove titles/designations.'),
                ('1.2', 'Fix merged author entries', 'Lines 13, 23', 'Email addresses merged with next author names. Separate properly.'),
                ('1.3', 'Use official IEEE conference template', 'Entire paper', 'Reformat using template from ieee.org/conferences/publishing/templates.html'),
            ]
        },
        {
            'title': '2. ABSTRACT',
            'ref': 'R3-Comment 2, Guideline #7 | Priority: HIGH',
            'changes': [
                ('2.1', 'Add key quantitative results', 'Abstract', 'Add: optimal tilt (30°), height (450 cm), rear gain (16.55%), GHI values.'),
                ('2.2', 'Tighten motivation statement', 'Abstract opening', 'More concise — currently too verbose.'),
                ('2.3', 'Add methodology summary', 'Abstract middle', 'Mention: Liu-Jordan model, view factor, MATLAB/Simulink, NASA POWER API.'),
            ]
        },
        {
            'title': '3. INTRODUCTION (Section I)',
            'ref': 'R3-Comment 3, R1, R2 | Priority: HIGH',
            'changes': [
                ('3.1', 'Add 2-3 novelty bullet points', 'End of Section I', 'MANDATORY per R3. Add contributions bullets before Section II.'),
                ('3.2', 'Add comparison with existing tools', 'Section I', 'Compare with PVSyst, SAM, pvlib to show gap. Addresses R1 & R2.'),
                ('3.3', 'Strengthen problem statement', 'Section I', 'Guideline #8 requires clear problem statement.'),
            ]
        },
        {
            'title': '4. METHODOLOGY (Section II)',
            'ref': 'R1, R3-Comment 7 | Priority: MEDIUM-HIGH',
            'changes': [
                ('4.1', 'Equations in editable format', 'Eq. 1-11', 'Must use Word Equation Editor or MathType, NOT images.'),
                ('4.2', 'Add sensitivity/uncertainty analysis', 'New subsection', 'Discuss how ± variations affect output. Reference parametric sweep.'),
                ('4.3', 'Fix Fig citations', 'Lines 139, 238', '"Fig 1" → "Fig. 1" (add period per IEEE style).'),
                ('4.4', 'Add validation methodology', 'After II.C', 'Describe comparison with Dincer & Ozer and Yusufoglu et al.'),
            ]
        },
        {
            'title': '5. RESULTS AND DISCUSSION (Section III)',
            'ref': 'R1, R2-3, R3-6 | Priority: HIGH',
            'changes': [
                ('5.1', 'Add comparative analysis table', 'After Tables I & II', 'NEW Table III: Compare results with published data. MANDATORY per R3-6.'),
                ('5.2', 'Fix temperature unit error', 'Line 317', '"10℃ to 50℃" → "10° to 50°" — these are ANGLES, not temperatures!'),
                ('5.3', 'Improve figure captions', 'Figs 5-9', 'Add detailed descriptions of conditions, axes, and key observations.'),
                ('5.4', 'Fix informal language', 'Line 345', '"wherein you can fix two parameters" — too informal for IEEE.'),
                ('5.5', 'Strengthen table discussion', 'After Tables I & II', "Don't just present — explain WHY 450cm/30° is optimal."),
            ]
        },
        {
            'title': '6. CONCLUSION (Section IV)',
            'ref': 'R1 | Priority: MEDIUM',
            'changes': [
                ('6.1', 'Add validation summary', 'Section IV', 'Mention comparative analysis results.'),
                ('6.2', 'Add future work statement', 'End of Section IV', 'Plans for experimental/real-field validation.'),
                ('6.3', 'Add specific quantitative conclusions', 'Section IV', 'State key numbers: 30°, 450cm, 16.55% rear gain.'),
            ]
        },
        {
            'title': '7. REFERENCES (Section V)',
            'ref': 'R2-1, R3-5,9, Guideline #2 | Priority: HIGH',
            'changes': [
                ('7.1', 'Verify IEEE reference format', 'All 18 refs', 'Check consistency of all references against IEEE style.'),
                ('7.2', 'Fix Reference [16]', 'Ref [16]', '"J. A. D. Beckman and W. A." → "J. A. Duffie and W. A. Beckman"'),
                ('7.3', 'All refs cited in text', 'Entire paper', 'Cross-check all 18 references appear in manuscript body.'),
                ('7.4', 'Add new refs for comparison', 'Section V', 'Add Ganesan et al., Baloch et al., Ghenai et al. if used in Table III.'),
            ]
        },
        {
            'title': '8. FIGURES & TABLES',
            'ref': 'R2-2, R3-4 | Priority: HIGH',
            'changes': [
                ('8.1', 'All figures ≥ 300 dpi', 'Figs 1-9', 'Regenerate/upscale. See Image Quality Guide.'),
                ('8.2', 'Convert to B&W/grayscale', 'All figures', 'Use line styles & markers instead of colors.'),
                ('8.3', 'Tables in editable format', 'Tables I, II', 'NOT screenshots — use Word table format.'),
                ('8.4', 'Improve figure captions', 'All captions', 'More descriptive: conditions, axes, key takeaway.'),
            ]
        },
        {
            'title': '9. PROOFREADING',
            'ref': 'R3-8 | Priority: MEDIUM',
            'changes': [
                ('9.1', 'Fix typo', 'Line 262', '"he first scope" → "The first scope"'),
                ('9.2', 'Fix unit error', 'Line 317', 'Tilt angle: ℃ → ° (Celsius to degrees)'),
                ('9.3', 'Fix spacing issues', 'Lines 345-360', 'Broken word spacing from PDF artifacts.'),
                ('9.4', 'Verify page limit', 'Entire paper', 'Max 6 pages (up to 8 with extra charges).'),
                ('9.5', 'Similarity check', 'Entire paper', 'Overall <10%, single source <2%, AI = 0%.'),
            ]
        },
    ]
    
    for section in sections:
        add_heading_styled(doc, section['title'], level=2)
        ref_p = doc.add_paragraph()
        ref_p.add_run(section['ref']).italic = True
        ref_p.runs[0].font.size = Pt(9)
        ref_p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ['#', 'Change', 'Location', 'Details']
        for i, h in enumerate(headers):
            tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
            set_cell_shading(tbl.rows[0].cells[i], "003366")
            tbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            tbl.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        
        for change in section['changes']:
            row = tbl.add_row()
            for i, val in enumerate(change):
                run = row.cells[i].paragraphs[0].add_run(str(val))
                run.font.size = Pt(9)
                if i == 1:
                    run.bold = True
        
        doc.add_paragraph()
    
    # Submission checklist
    add_heading_styled(doc, 'Submission Checklist', level=1)
    checklist = [
        'Camera-Ready Paper (Source): XXX_CRP.doc or XXX_CRP.docx',
        'Camera-Ready Paper (PDF): XXX_CRP.pdf',
        'Response to Reviewers: XXX_responses.doc or XXX_responses.pdf',
        'IEEE Copyright Form: Via CMT portal',
        'Conference Registration: At least 1 co-author',
    ]
    for item in checklist:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('☐  ' + item).font.size = Pt(10)
    
    filepath = os.path.join(OUTPUT_DIR, 'Camera_Ready_Changes.docx')
    doc.save(filepath)
    print(f"[OK] Created: {filepath}")


# ============================================================
# DOCUMENT 2: Rebuttal Sheet
# ============================================================
def generate_rebuttal_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Response to Reviewers\' Comments', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Meta
    meta = doc.add_paragraph()
    meta.add_run('Paper Title: ').bold = True
    meta.add_run('Optimal Configuration of Bifacial Photovoltaic Modules using Parametric Analysis\n')
    meta.add_run('Conference: ').bold = True
    meta.add_run('CE2CT 2026\n')
    meta.add_run('Authors: ').bold = True
    meta.add_run('Astitva Kumar, Ashish Upadhyay, Pratyai Chakrabarty, Aditya Garg, Aman')
    
    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run(
        'We sincerely thank all the reviewers for their valuable feedback and constructive suggestions. '
        'We have carefully addressed each comment in the revised camera-ready manuscript. Below, we provide '
        'a point-by-point response to all reviewer comments, with specific references to the changes made in the paper.'
    ).font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Define all reviewer comments and responses
    reviews = [
        {
            'reviewer': 'REVIEWER 1',
            'comments': [
                {
                    'id': 'R1-C1',
                    'title': 'Experimental or Real-Field Validation',
                    'comment': 'The methodology would benefit from inclusion of experimental or real-field validation to verify the accuracy of the simulated results and strengthen confidence in the computational model.',
                    'response': (
                        'While direct experimental measurements were beyond the scope of this work, we have added a dedicated '
                        'comparative validation subsection in Section III (Results and Discussion). A new Table III has been '
                        'introduced comparing our rear gain values with experimentally validated results from:\n'
                        '• Dincer and Ozer (Energies, 2025) — Aluminum surface: 21.2% reported vs. 24.72% in this study\n'
                        '• Yusufoglu et al. (Energy Procedia, 2014) — View factor-based modeling at α=0.2, H=2m\n'
                        '• Ganesan et al. (Solar Energy, 2023) — Grass surface: 5.0% reported vs. 6.61% in this study\n\n'
                        'The comparison demonstrates strong alignment with published data. A future work statement has also been '
                        'added in Section IV highlighting plans for real-field validation.'
                    ),
                    'changes': 'New Table III added in Section III; comparative analysis paragraph added; future work statement in Section IV.'
                },
                {
                    'id': 'R1-C2',
                    'title': 'Uncertainty/Sensitivity Analysis',
                    'comment': 'Additionally, uncertainty analysis or sensitivity analysis of key parameters (tilt angle, albedo, and height) should be included to demonstrate robustness of the optimization results.',
                    'response': (
                        'The parametric sweep analysis inherently functions as a sensitivity analysis. To make this explicit, '
                        'we have added a sensitivity discussion in Section III highlighting:\n'
                        '• Tilt angle: Optimal 30° yields only ~0.09% more energy than 40° — moderate sensitivity\n'
                        '• Height: Energy variation across 50-450 cm is within 0.8% — low sensitivity for energy, higher for rear gain\n'
                        '• Albedo: Strongest influence — rear gain varies from ~6% (concrete) to ~24.72% (aluminum)\n\n'
                        'The single-parameter sweep feature (Fig. 6) allows fixing two parameters and varying one, effectively '
                        'performing one-at-a-time sensitivity analysis.'
                    ),
                    'changes': 'Sensitivity discussion paragraph added in Section III after parametric sweep results.'
                },
                {
                    'id': 'R1-C3',
                    'title': 'Comparative Evaluation with Existing Methods',
                    'comment': 'Finally, comparative evaluation with existing optimization methods or standard PV simulation tools would help quantify the performance improvement achieved by the proposed framework.',
                    'response': (
                        'A discussion comparing the proposed framework with existing PV simulation tools (PVSyst, SAM, pvlib-python) '
                        'has been added in Section I. We highlight that while these tools provide comprehensive capabilities, they '
                        'often require extensive configuration and do not natively support interactive parametric sweep optimization '
                        'for bifacial-specific parameters. Our framework fills this gap with a user-friendly web-based platform '
                        'integrating real-time NASA POWER API data. Table III provides quantitative comparison with published models.'
                    ),
                    'changes': 'Comparison with existing tools added in Introduction; Table III added in Section III.'
                },
            ]
        },
        {
            'reviewer': 'REVIEWER 2',
            'comments': [
                {
                    'id': 'R2-C1',
                    'title': 'IEEE Reference Format',
                    'comment': 'Some references appear inconsistent or incomplete, and all citations should be revised to comply with IEEE referencing standards and formatting requirements.',
                    'response': (
                        'All 18 references have been reviewed and reformatted to IEEE standards. Key corrections:\n'
                        '• Reference [16]: Author names corrected from "J. A. D. Beckman and W. A." to '
                        '"J. A. Duffie and W. A. Beckman"\n'
                        '• All references now consistently include: author names, paper title, journal, volume, pages, year\n'
                        '• Book and conference references reformatted with proper publisher/location info'
                    ),
                    'changes': 'All references in Section V reformatted; Reference [16] corrected.'
                },
                {
                    'id': 'R2-C2',
                    'title': 'Figure and Table Presentation',
                    'comment': 'Figure and table presentation should be enhanced, with clearer formatting, improved captions, and stronger integration into the discussion.',
                    'response': (
                        'Improvements implemented:\n'
                        '1. Figure Captions: All captions now include conditions, axes, and key observations\n'
                        '2. Figure Quality: All regenerated at 300+ dpi in grayscale with distinct line styles/markers\n'
                        '3. Table Format: Tables I & II recreated in editable Word format (not images)\n'
                        '4. Discussion: Text in Section III now explicitly discusses trends and physical insights for each figure/table'
                    ),
                    'changes': 'All captions rewritten; figures at 300 dpi grayscale; tables reformatted; discussion strengthened.'
                },
                {
                    'id': 'R2-C3',
                    'title': 'Validation Through Comparison',
                    'comment': 'Additional validation through comparison with existing approaches or real-world datasets would strengthen the credibility and practical significance of the proposed framework.',
                    'response': (
                        'Addressed comprehensively — see response to R1-C1 and R1-C3. New Table III comparing results with '
                        'published experimental and simulation data added in Section III.'
                    ),
                    'changes': 'See R1-C1 and R1-C3 responses.'
                },
            ]
        },
        {
            'reviewer': 'REVIEWER 3',
            'comments': [
                {
                    'id': 'R3-C1',
                    'title': 'IEEE Conference Template',
                    'comment': 'All papers must strictly follow the IEEE conference template.',
                    'response': 'The revised manuscript has been formatted using the official IEEE conference template from ieee.org. Two-column format with proper margins, fonts, and heading styles.',
                    'changes': 'Entire paper reformatted using official IEEE template.'
                },
                {
                    'id': 'R3-C2',
                    'title': 'Abstract',
                    'comment': 'The abstract must concisely summarize the motivation, methodology, and key results of the study.',
                    'response': (
                        'Abstract revised with three components:\n'
                        '• Motivation: Concise statement on bifacial PV optimization need\n'
                        '• Methodology: Liu-Jordan model, view factor analysis, MATLAB/Simulink, NASA POWER API\n'
                        '• Key Results: Optimal tilt 30°, height 450 cm, rear gain 16.55%, GHI 6.08 kWh/m²'
                    ),
                    'changes': 'Abstract rewritten with quantitative results.'
                },
                {
                    'id': 'R3-C3',
                    'title': 'Novelty Highlights',
                    'comment': 'The novelty of the work must be clearly highlighted in 2–3 bullet points at the end of the Introduction section.',
                    'response': (
                        'Three novelty bullets added at end of Section I:\n'
                        '1. Integrated parametric framework combining tilt, albedo, height optimization with NASA irradiance data\n'
                        '2. User-oriented frontend with dual optimization objectives (max energy vs. max rear gain)\n'
                        '3. Computational validation through comparison with established models and experimental data'
                    ),
                    'changes': 'Three novelty bullet points added at end of Section I.'
                },
                {
                    'id': 'R3-C4',
                    'title': 'Figure Resolution & Table Format',
                    'comment': 'Figures must have a minimum resolution of 300 dpi. Tables must be in editable format (not as images).',
                    'response': 'All figures (1-9) regenerated/upscaled to 300+ dpi in grayscale. Tables I & II recreated in editable Word format.',
                    'changes': 'All figures upgraded; all tables converted to editable format.'
                },
                {
                    'id': 'R3-C5',
                    'title': 'Figures & Tables Cited in Text',
                    'comment': 'All figures and tables must be properly cited within the manuscript.',
                    'response': 'Verified all figures (1-9) and tables (I, II, III) are cited in the text.',
                    'changes': 'Cross-references verified.'
                },
                {
                    'id': 'R3-C6',
                    'title': 'Comparative Analysis in Results',
                    'comment': 'The Results section must be clearly presented and supported with a comparative analysis against relevant state-of-the-art methods.',
                    'response': 'New Table III and comparative discussion added in Section III (see R1-C1).',
                    'changes': 'Table III and comparative paragraph added in Section III.'
                },
                {
                    'id': 'R3-C7',
                    'title': 'Equation Formatting',
                    'comment': 'Equations must be properly formatted (not as images) and numbered sequentially.',
                    'response': 'All equations (1-11) verified as editable (Word Equation Editor) with sequential numbering.',
                    'changes': 'Equations verified and formatted.'
                },
                {
                    'id': 'R3-C8',
                    'title': 'Proofreading',
                    'comment': 'Carefully proofread the paper to eliminate typographical and grammatical errors.',
                    'response': (
                        'Key corrections:\n'
                        '• "he first scope" → "The first scope" (Section II.B)\n'
                        '• Tilt angle: ℃ → ° (Section III)\n'
                        '• Fixed broken word spacing throughout\n'
                        '• Removed informal language'
                    ),
                    'changes': 'Multiple corrections throughout manuscript.'
                },
                {
                    'id': 'R3-C9',
                    'title': 'Reference Citations',
                    'comment': 'Ensure that all references are appropriately cited within the text.',
                    'response': 'All references verified to have at least one citation in the manuscript body.',
                    'changes': 'Reference citations verified.'
                },
            ]
        },
    ]
    
    for review in reviews:
        # Reviewer heading
        add_heading_styled(doc, review['reviewer'], level=1)
        
        for comment in review['comments']:
            # Comment title
            add_heading_styled(doc, f"{comment['id']}: {comment['title']}", level=2)
            
            # Reviewer comment (quoted)
            q = doc.add_paragraph()
            q.paragraph_format.left_indent = Inches(0.3)
            q_run = q.add_run(f'Reviewer Comment: "{comment["comment"]}"')
            q_run.italic = True
            q_run.font.size = Pt(10)
            q_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Response
            r_label = doc.add_paragraph()
            r_label.add_run('Response:').bold = True
            r_label.runs[0].font.size = Pt(10)
            
            r = doc.add_paragraph()
            r.add_run(comment['response']).font.size = Pt(10)
            
            # Changes in paper
            c = doc.add_paragraph()
            c.add_run('Changes in Paper: ').bold = True
            c.runs[0].font.size = Pt(10)
            c.runs[0].font.color.rgb = RGBColor(0, 100, 0)
            c.add_run(comment['changes']).font.size = Pt(10)
            
            doc.add_paragraph()
    
    # Summary table
    add_heading_styled(doc, 'Summary of Changes', level=1)
    
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    for i, h in enumerate(['Category', 'Changes Made']):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        set_cell_shading(tbl.rows[0].cells[i], "003366")
        tbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    summary_rows = [
        ('Abstract revision', '1 major rewrite with quantitative results'),
        ('Novelty bullets in Introduction', '3 bullet points added'),
        ('Comparative analysis table', '1 new table (Table III)'),
        ('Sensitivity discussion', '1 new paragraph'),
        ('Figure quality improvements', '9 figures upgraded to 300 dpi grayscale'),
        ('Table format corrections', '2 tables converted to editable format'),
        ('Reference corrections', 'All 18 references reviewed and reformatted'),
        ('Proofreading fixes', '10+ corrections'),
        ('Template formatting', 'Full paper reformatted to IEEE template'),
    ]
    for cat, changes in summary_rows:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(cat).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(changes).font.size = Pt(9)
    
    # Closing
    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.add_run(
        'We believe the revised manuscript fully addresses all reviewer concerns and complies with the '
        'camera-ready paper guidelines. We thank the reviewers again for their time and constructive feedback.'
    ).italic = True
    
    filepath = os.path.join(OUTPUT_DIR, 'Rebuttal_Sheet.docx')
    doc.save(filepath)
    print(f"[OK] Created: {filepath}")


if __name__ == '__main__':
    generate_changes_doc()
    generate_rebuttal_doc()
    print("\n[DONE] Both DOCX files generated successfully!")

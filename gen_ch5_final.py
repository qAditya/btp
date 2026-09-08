"""
Generate Chapter 5 + References + Appendices (with actual code) as standalone DOCX.
Formatting: Times New Roman 12pt body / 14pt headings, 1.5 spacing, margins L=1.5" R=T=B=1"
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH as A
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document()
for s in doc.sections:
    s.top_margin=Inches(1);s.bottom_margin=Inches(1);s.left_margin=Inches(1.5);s.right_margin=Inches(1)
st=doc.styles['Normal'];st.font.name='Times New Roman';st.font.size=Pt(12);st.paragraph_format.line_spacing=1.5

def H(t,sz=14,al=A.LEFT):
    p=doc.add_paragraph();p.alignment=al;p.paragraph_format.space_before=Pt(12);p.paragraph_format.space_after=Pt(6)
    r=p.add_run(t);r.bold=True;r.font.size=Pt(sz);r.font.name='Times New Roman'
def P(t,al=A.JUSTIFY):
    p=doc.add_paragraph();p.alignment=al;p.paragraph_format.space_after=Pt(6);p.paragraph_format.space_before=Pt(3)
    r=p.add_run(t);r.font.size=Pt(12);r.font.name='Times New Roman'
def NL():
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(0);p.paragraph_format.space_after=Pt(0)
def CODE(title, code_text):
    """Add a code block with monospaced font."""
    H(title, 12)
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(8)

# =====================================================================
# CHAPTER 5
# =====================================================================
H("CHAPTER 5", 16, A.CENTER)
H("CONCLUSION AND FUTURE WORK", 14, A.CENTER)
NL()

H("5.1  Summary of Contributions")
P("This project has presented a comprehensive computational framework for the parametric analysis and optimization of bifacial photovoltaic systems. The work addresses the identified gaps in existing literature by providing a systematic, accessible, and location-specific tool for determining optimal installation configurations for bifacial PV modules. The key contributions of this work are summarized as follows.")

P("First, a rigorous mathematical model for computing the front-side and rear-side irradiance on bifacial PV modules has been implemented. The model incorporates the Liu-Jordan isotropic sky model [1] for front-side irradiance estimation, the Erbs clearness-index correlation for diffuse-direct decomposition, and the Yusufoglu view factor approach [2] for accurate rear-side irradiance calculation with shadow correction. The formulation accounts for the combined effects of tilt angle, ground albedo, mounting height, solar geometry, and panel azimuth on both front-side and rear-side irradiance contributions, providing a physically consistent representation of bifacial PV module operation under realistic conditions.")

P("Second, a MATLAB/Simulink model has been developed for simulating the electrical characteristics of bifacial PV modules. The Simulink model utilizes a single-diode equivalent circuit representation of the PV module and incorporates the NOCT-based thermal model for temperature derating. The model enables the computation of complete I-V and P-V characteristic curves under any combination of irradiance and temperature conditions, providing validated electrical output predictions for system design and performance assessment.")

P("Third, a user-oriented frontend platform has been designed and implemented that integrates location-based solar irradiance data retrieval from the NASA POWER API with automated parametric sweep analyses. The platform features a modular architecture comprising a responsive web-based user interface, a Node.js/Express backend application layer, and dual analysis engines (JavaScript-based and MATLAB-based). The platform provides an accessible tool for PV system designers that requires no specialized software knowledge or access to commercial simulation tools, directly addressing the accessibility gap identified in the literature review.")

P("Fourth, the framework supports dual optimization objectives, namely maximum total energy generation and maximum rear-side gain. This dual optimization capability enables system designers to evaluate the trade-offs between these complementary objectives and make informed configuration decisions based on their specific project requirements and priorities. Few existing tools in the literature provide this level of optimization flexibility for bifacial PV systems.")

P("Fifth, the proposed approach has been validated through a detailed case study for Greater Noida, India (28.47\u00b0N, 77.50\u00b0E), with results demonstrating good quantitative agreement with published literature including the parametric analysis of Dincer and Ozer [3] and the simulation study of Yusufoglu et al. [2]. The validation confirms the reliability and accuracy of both the mathematical formulation and its computational implementation across diverse installation configurations.")

H("5.2  Key Findings")
P("The parametric analyses conducted in this study have yielded the following key findings that provide practical guidance for the design and optimization of bifacial PV installations:")

P("(i) Optimal Configuration for Maximum Energy: For the case study location (Greater Noida, India) with a concrete ground surface (albedo = 0.3), the optimal configuration for maximum energy output was identified as a tilt angle of 30 degrees and mounting height of 450 cm, achieving a total daily energy of 3.168 kWh per module and a rear gain of 16.55%. The optimal tilt angle of 30 degrees closely corresponds to the site latitude (28.47\u00b0N), consistent with established PV design principles for annual energy maximization.")

P("(ii) Optimal Configuration for Maximum Rear Gain: The optimal configuration for maximum rear gain was identified as a tilt angle of 10 degrees and height of 450 cm, achieving a rear gain of 17.90% but with a lower total energy of 2.994 kWh (approximately 5.5% lower than the maximum energy configuration). This finding demonstrates the existence of a clear and quantifiable trade-off between total energy maximization and rear gain maximization that must be considered during system design.")

P("(iii) Dominant Influence of Albedo: Ground surface albedo was confirmed as the most influential individual parameter affecting rear-side irradiance contribution. The rear gain varied from below 5% for low-albedo surfaces (such as dry asphalt with albedo approximately 0.12) to over 24% for highly reflective surfaces (such as aluminum with albedo 0.85), representing a nearly five-fold variation. The relationship between albedo and rear gain was found to be approximately linear, consistent with the mathematical formulation in equation (10) and independently confirmed by the simulation results of Yusufoglu et al. [2]. This finding has direct practical implications for site preparation and ground surface selection in bifacial PV installations.")

P("(iv) Logarithmic Height Dependence: Mounting height exhibited a positive but diminishing effect on rear-side performance, consistent with a near-logarithmic dependence reported by Yusufoglu et al. [2]. The practical implication is that mounting heights in the range of 200-300 cm provide a good balance between performance improvement and structural feasibility for most ground-mounted applications. Beyond this range, the marginal energy benefit becomes insufficient to justify the additional structural cost and complexity.")

P("(v) Necessity of Combined Optimization: The interaction between tilt angle, height, and albedo necessitates combined parametric optimization rather than independent single-parameter studies. The combined sweep approach captured interaction effects that would be missed by analyzing parameters in isolation. Furthermore, the system performance was found to be relatively robust to small variations near the optimal configuration, with the top five configurations differing by less than 0.25% in total energy, providing practical flexibility for accommodating site-specific constraints.")

P("(vi) Consistency with Published Literature: The comparative analysis with published studies confirmed that the proposed framework produces results that are quantitatively consistent with established research. The rear gain values, albedo sensitivity trends, height dependency patterns, and tilt angle optimization results all showed good agreement with the reference studies of Dincer and Ozer [3], Yusufoglu et al. [2], Ganesan et al. [4], and Pelaez et al. [5], validating the accuracy and reliability of the implemented mathematical models.")

H("5.3  Limitations")
P("The present work has certain limitations that should be acknowledged and considered when interpreting the results and applying the framework to real-world system design:")

P("(i) Uniform Albedo Assumption: The mathematical model assumes uniform ground albedo beneath and surrounding the module. In real-world installations, ground surface properties may vary spatially due to vegetation growth, shadows from nearby structures, mixed surface types, or degradation of artificial reflective surfaces over time. Non-uniform albedo distributions can lead to spatially varying rear-side irradiance that is not captured by the present model.")

P("(ii) Single-Day Analysis: The case study analysis was performed using single-day irradiance data from the NASA POWER API. While this approach provides useful insights into the relative performance of different configurations and enables rapid optimization, extending the analysis to annual energy yield computations incorporating seasonal variations in solar geometry, temperature, and irradiance conditions would provide more comprehensive and robust optimization guidance for long-term system design decisions.")

P("(iii) Single-Row Configuration: The current implementation models a single isolated bifacial PV module and does not account for inter-row shading effects that occur in multi-row array configurations. In large-scale solar installations, adjacent rows cast shadows on the ground between them, reducing the ground-reflected radiation available to the rear surface and potentially altering the optimal configuration parameters.")

P("(iv) Simplified Thermal Model: Temperature effects on module performance are modeled using the simplified NOCT approach, which relates cell temperature to ambient temperature and incident irradiance through a linear approximation. More detailed thermal models that account for wind speed, convective cooling, mounting configuration effects on rear-surface ventilation, and radiative heat exchange may improve the accuracy of energy yield predictions, particularly in hot climatic conditions.")

P("(v) Isotropic Sky Model: The Liu-Jordan isotropic sky model used for diffuse radiation estimation assumes uniform distribution of diffuse radiation across the sky hemisphere. This assumption may underestimate the circumsolar and horizon brightening components that contribute to tilted surface irradiance, particularly under partly cloudy conditions. More advanced anisotropic models such as the Perez model could improve the accuracy of front-side irradiance estimation.")

P("(vi) Spectral Effects: The current model does not account for the spectral dependence of ground surface reflectivity or the wavelength-dependent response of bifacial silicon solar cells. As demonstrated by Riedel-Lyngskjaer et al. [6], spectral albedo effects can influence the actual electrical output from ground-reflected radiation, particularly for surfaces with strong spectral selectivity.")

H("5.4  Scope for Future Work")
P("Several promising directions for future research and development are identified based on the outcomes, insights, and limitations of this project:")

P("(i) Extension to Annual Energy Yield: The current framework can be extended to perform year-round energy yield simulations by aggregating daily analyses across all days of the year. This extension would incorporate seasonal variations in solar geometry (declination, day length), ambient temperature profiles, and cloud cover patterns, enabling more robust design decisions that account for the full range of operating conditions experienced throughout the year. Annual yield predictions are essential for accurate financial analysis and investment decision-making in commercial PV projects.")

P("(ii) Multi-Row Array Modeling: Incorporating inter-row shading models would enable the framework to optimize row-to-row spacing as an additional design variable alongside tilt angle, height, and albedo. This enhancement is particularly important for utility-scale bifacial solar farm design, where inter-row interactions significantly affect system performance and where ground coverage ratio optimization has substantial economic implications.")

P("(iii) Machine Learning-Based Rapid Optimization: Developing machine learning models (such as neural networks, random forests, or gradient boosting machines) trained on comprehensive parametric sweep databases could enable near-instantaneous prediction of optimal configurations for new locations without requiring full computational analysis. Such surrogate models could dramatically reduce the optimization time while maintaining acceptable accuracy for preliminary system design.")

P("(iv) Techno-Economic Analysis Integration: Integrating economic analysis capabilities would allow the framework to evaluate the cost-effectiveness of different bifacial PV configurations through levelized cost of energy (LCOE) calculations. This analysis would incorporate structural costs associated with different mounting heights, ground preparation costs for albedo enhancement measures, and the economic value of additional energy generation, enabling more holistic design optimization that considers both technical performance and financial viability.")

P("(v) Real-Time Performance Monitoring: Developing a real-time monitoring and analytics system that continuously compares actual system performance with predicted values could enable adaptive optimization strategies and early detection of performance degradation. Integration with IoT sensors, edge computing devices, and cloud-based analytics platforms could provide automated performance assessment, fault detection, and predictive maintenance alerts.")

P("(vi) Geographic Information System Integration: Combining the optimization framework with GIS tools and remote sensing data could enable automated site assessment for large-scale bifacial PV deployment planning. This integration would incorporate terrain elevation models, land use classification, local albedo mapping from satellite imagery, and proximity analysis to grid infrastructure, providing comprehensive site suitability analysis at regional and national scales.")

P("(vii) Tracking System Analysis: Extending the framework to support single-axis and dual-axis tracking configurations would enable direct comparison of fixed-tilt and tracking bifacial PV systems under identical environmental conditions. Tracking systems can significantly enhance bifacial module performance by maintaining optimal front-side irradiance throughout the day while simultaneously affecting the shadow geometry and rear-side irradiance patterns.")

P("(viii) Advanced Irradiance Models: Implementing anisotropic sky models (such as the Perez or Hay-Davies model) and spectral irradiance decomposition would improve the accuracy of both front-side and rear-side irradiance estimates, particularly for locations with complex sky conditions or for ground surfaces with strong spectral selectivity.")

H("5.5  Task, Achievements, and Possible Beneficiaries")
P("The primary task of this project was to develop and validate a computational framework for parametric analysis and optimization of bifacial PV installations using location-specific solar irradiance data. This task has been successfully accomplished through the design and implementation of an integrated platform comprising rigorous mathematical modeling, MATLAB/Simulink electrical simulation, and a web-based frontend interface with NASA POWER API integration.")

P("The key achievements of this project include: (a) development and implementation of a validated mathematical model for bifacial PV irradiance computation incorporating view factor geometry and shadow effects; (b) creation of a MATLAB/Simulink model for electrical performance simulation of bifacial PV modules; (c) implementation of a user-friendly, web-based frontend platform with automated parametric sweep analysis capabilities; (d) demonstration of dual optimization (maximum energy and maximum rear gain) through a comprehensive case study; and (e) acceptance of a research paper based on this work at the IC2PCT 2026 International Conference.")

P("The possible beneficiaries of this work include: PV system designers and engineers seeking optimal configuration guidance for bifacial installations; solar project developers evaluating the feasibility and expected performance enhancement of bifacial PV technology; EPC (Engineering, Procurement, and Construction) contractors requiring quick, reliable optimization guidance for specific project sites; academic researchers studying bifacial PV technology who need a flexible, validated simulation platform; and policymakers and energy planners assessing the potential of bifacial PV technology for meeting renewable energy targets.")

doc.add_page_break()

# =====================================================================
# REFERENCES (IEEE format, citation order matching the report)
# =====================================================================
H("REFERENCES", 16, A.CENTER)
NL()

refs = [
'[1] B. Y. H. Liu and R. C. Jordan, "The interrelationship and characteristic distribution of direct, diffuse and total solar radiation," Solar Energy, vol. 4, no. 3, pp. 1-19, 1960.',
'[2] U. A. Yusufoglu, T. M. Pletzer, L. J. Koduvelikulathu, C. Comparotto, R. Kopecek, and H. Kurz, "Analysis of the annual performance of bifacial modules and optimization methods," IEEE J. Photovolt., vol. 5, no. 1, pp. 320-328, Jan. 2015.',
'[3] F. Dincer and E. Ozer, "Optimization of rear-side energy contribution in bifacial PV panels: A parametric analysis on albedo, tilt, height, and mounting configuration," Energies, vol. 18, no. 17, Art. no. 4443, 2025.',
'[4] K. Ganesan, D. P. Winston, S. Sugumar, and S. Jegan, "Performance analysis of n-type PERT bifacial solar PV module under diverse albedo conditions," Solar Energy, vol. 252, pp. 81-90, 2023.',
'[5] S. A. Pelaez, C. Deline, S. M. MacAlpine, B. Marion, J. S. Stein, and R. K. Kostuk, "Comparison of bifacial solar irradiance model predictions with field validation," IEEE J. Photovolt., vol. 9, no. 1, pp. 82-88, Jan. 2019.',
'[6] N. Riedel-Lyngskjaer, J. M. L. P. Larsen, and B. G. Nielsen, "Effect of spectral albedo on energy yield of bifacial photovoltaic systems," Solar Energy, vol. 231, pp. 176-186, 2022.',
'[7] R. Guerrero-Lemus, R. Vega, T. Kim, A. Kimm, and L. E. Shephard, "Bifacial solar photovoltaics - A technology review," Renew. Sustain. Energy Rev., vol. 60, pp. 1533-1549, 2016.',
'[8] M. Alam, M. R. Khan, and C. Deline, "Performance comparison between bifacial and monofacial photovoltaic systems under varying environmental conditions," Sustain. Energy Technol. Assess., vol. 57, Art. no. 103210, 2023.',
'[9] C. Deline, S. A. Pelaez, S. MacAlpine, and B. Marion, "Bifacial photovoltaic system performance: Separating fact from fiction," Nat. Renew. Energy Lab. (NREL), Golden, CO, USA, Tech. Rep. NREL/CP-5K00-74090, 2019.',
'[10] X. Sun, M. R. Khan, C. Deline, and M. A. Alam, "Optimization and performance of bifacial solar modules: A global perspective," Appl. Energy, vol. 212, pp. 1601-1610, 2018.',
'[11] A. Basak, S. Das, and R. Banerjee, "Tilt angle optimization for bifacial photovoltaic modules for enhanced energy generation," Appl. Energy, vol. 345, Art. no. 121245, 2025.',
'[12] U. Peter and M. Novak, "Optimal tilt angle for maximizing energy production of bifacial solar panels," J. Sustain. Energy Syst., vol. 14, no. 1, pp. 45-56, 2025.',
'[13] M. H. Aksoy and H. A. Ceylan, "Investigation of ground albedo effects on bifacial photovoltaic panel performance," Renew. Energy Res. J., vol. 13, no. 4, pp. 1800-1808, 2023.',
'[14] N. Baghel, A. Chandel, and S. K. Sharma, "Performance evaluation and optimization of albedo and tilt angle in bifacial photovoltaic systems," Solar Energy, vol. 250, pp. 84-95, 2023.',
'[15] I. N. Atalay, M. Kaya, and H. Ozturk, "Experimental investigation of ground surface reflectivity impact on bifacial photovoltaic systems," Solar Energy, vol. 246, pp. 65-74, 2023.',
'[16] D. S. Braga, R. Zilles, and E. Lorenzo, "Performance analysis of bifacial photovoltaic modules under tropical climatic conditions," Renew. Energy Environ. Sustain., vol. 8, pp. 1-10, 2023.',
'[17] A. F. Almarshoud, A. Al-Badi, and M. Al-Sulaiman, "Experimental performance analysis of bifacial photovoltaic modules under different operating conditions," Energies, vol. 17, no. 21, Art. no. 5456, 2024.',
'[18] M. T. Patel, M. R. Khan, and M. A. Alam, "Optimum design of tracking bifacial solar farms," IEEE J. Photovolt., vol. 11, no. 1, pp. 237-245, Jan. 2021.',
'[19] J. A. Duffie and W. A. Beckman, Solar Engineering of Thermal Processes, 4th ed. Hoboken, NJ, USA: Wiley, 2013.',
'[20] B. Marion, S. MacAlpine, C. Deline, A. Asgharzadeh, F. Toor, D. Riley, J. Stein, and C. Hansen, "A practical irradiance model for bifacial PV modules," in Proc. 44th IEEE Photovolt. Spec. Conf. (PVSC), Washington, DC, USA, 2017, pp. 1537-1542.',
'[21] P. K. Sahu, J. N. Roy, and C. Chakraborty, "Performance assessment of a bifacial PV system using a new energy estimation model," Solar Energy, vol. 262, Art. no. 111818, 2023.',
'[22] C. Ghenai, F. F. Ahmad, O. Rejeb, and A. K. Hamid, "Sensitivity analysis of design parameters and power gain correlations of bi-facial solar PV system using response surface methodology," Solar Energy, vol. 223, pp. 44-53, 2021.',
]
for ref in refs:
    P(ref, al=A.LEFT)

doc.add_page_break()

# =====================================================================
# APPENDIX A: MATLAB Code
# =====================================================================
H("APPENDIX A", 16, A.CENTER)
H("MATLAB/Simulink Code", 14, A.CENTER)
NL()
P("This appendix contains the MATLAB code used for the bifacial PV irradiance computation and parametric sweep analysis. The code implements the mathematical formulation described in Chapter 3 of the report.")

# Read actual code files
calc_code = open(r'backend\simulink\functions\calculate_irradiance.m', 'r', encoding='utf-8').read()
sweep_code = open(r'backend\simulink\scripts\run_bifacial_sweep.m', 'r', encoding='utf-8').read()

CODE("A.1  Irradiance Calculation Function (calculate_irradiance.m)", calc_code)
doc.add_page_break()
CODE("A.2  Parametric Sweep Script (run_bifacial_sweep.m)", sweep_code)
doc.add_page_break()

# =====================================================================
# APPENDIX B: Backend Code
# =====================================================================
H("APPENDIX B", 16, A.CENTER)
H("Backend Application Code", 14, A.CENTER)
NL()
P("This appendix contains excerpts of the backend application code implemented in Node.js/Express for the NASA POWER API integration and simulation service.")

irr_code = open(r'backend\src\services\irradianceService.js', 'r', encoding='utf-8').read()
server_code = open(r'backend\src\server.js', 'r', encoding='utf-8').read()

CODE("B.1  Express Server Entry Point (server.js)", server_code)
NL()
CODE("B.2  NASA POWER API Integration Service (irradianceService.js)", irr_code)
doc.add_page_break()

# Read simulation service (key excerpts - first 250 lines for view factor + irradiance calc)
sim_code_full = open(r'backend\src\services\simulationService.js', 'r', encoding='utf-8').read()
# Extract key sections
sim_lines = sim_code_full.split('\n')
sim_excerpt = '\n'.join(sim_lines[:220])  # Solar geometry + shadow view factor + diffuse fraction

CODE("B.3  Simulation Engine - Solar Geometry and Shadow View Factor (simulationService.js, excerpt)", sim_excerpt)
doc.add_page_break()

sim_excerpt2 = '\n'.join(sim_lines[566:745])  # calculateEffectiveIrradiance + evaluateConfiguration
CODE("B.4  Simulation Engine - Effective Irradiance Calculation (simulationService.js, excerpt)", sim_excerpt2)
doc.add_page_break()

# =====================================================================
# LIST OF PUBLICATIONS
# =====================================================================
H("LIST OF PUBLICATIONS", 14, A.CENTER)
NL()
P("The following publication has been produced as part of this project work:")
NL()
P('[1] A. Upadhyay, P. Chakrabarty, Aman, and A. Garg, "Performance analysis of bifacial PV system based on different albedos," accepted for presentation at IC2PCT 2026 (International Conference on Clean and Prospective Computing Technologies), 2026.', al=A.LEFT)
doc.add_page_break()

# =====================================================================
# BIO-DATA
# =====================================================================
H("BIO-DATA", 16, A.CENTER)
NL()
studs = [("Ashish Upadhyay","2022UEE4521"),("Pratyai Chakrabarty","2022UEE4586"),("Aman","2022UEE4532"),("Aditya Garg","2022UEE4503")]
for name, roll in studs:
    H(name, 12)
    P(f"Roll Number: {roll}")
    P("Programme: B.Tech. Electrical Engineering")
    P("Department: Electrical Engineering")
    P("Institution: Netaji Subhas University of Technology (NSUT), Dwarka, New Delhi - 110078")
    P("Batch: 2022-2026")
    NL()

# =====================================================================
# SAVE
# =====================================================================
doc.save('Chapter5_Refs_Appendices.docx')
print("=== SAVED: Chapter5_Refs_Appendices.docx ===")

# Count stats
wc = sum(len(p.text.split()) for p in doc.paragraphs)
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Word count: {wc}")
print(f"Estimated pages: ~{max(wc//250, 20)}")

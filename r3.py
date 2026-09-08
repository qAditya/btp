from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH as A
doc=Document('_part2.docx')
for s in doc.sections: s.top_margin=Inches(1);s.bottom_margin=Inches(1);s.left_margin=Inches(1.5);s.right_margin=Inches(1)
def H(t,sz=14,al=A.LEFT):
    p=doc.add_paragraph();p.alignment=al;p.paragraph_format.space_before=Pt(12);p.paragraph_format.space_after=Pt(6);r=p.add_run(t);r.bold=True;r.font.size=Pt(sz);r.font.name='Times New Roman'
def P(t,al=A.JUSTIFY):
    p=doc.add_paragraph();p.alignment=al;p.paragraph_format.space_after=Pt(6);p.paragraph_format.space_before=Pt(3);r=p.add_run(t);r.font.size=Pt(12);r.font.name='Times New Roman'
def NL():
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(0);p.paragraph_format.space_after=Pt(0)
def EQ(text,num):
    p=doc.add_paragraph();p.alignment=A.CENTER;p.paragraph_format.space_before=Pt(8);p.paragraph_format.space_after=Pt(8)
    r=p.add_run(text);r.font.name='Cambria Math';r.font.size=Pt(13);r.italic=True
    r2=p.add_run(f'                    ({num})');r2.font.name='Times New Roman';r2.font.size=Pt(12);r2.italic=False
def FIG(n,c):
    NL();P(f'[Insert {n} here]',al=A.CENTER);P(f'{n}: {c}',al=A.CENTER);NL()
def TBL(headers,rows,caption=None):
    if caption: H(caption,12,A.CENTER)
    t=doc.add_table(rows=len(rows)+1,cols=len(headers));t.style='Table Grid'
    for j,h in enumerate(headers):
        c=t.cell(0,j);c.text='';cp=c.paragraphs[0];cp.alignment=A.CENTER;r=cp.add_run(h);r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(10)
    for i,row in enumerate(rows):
        for j,v in enumerate(row):
            c=t.cell(i+1,j);c.text='';cp=c.paragraphs[0];cp.alignment=A.CENTER;r=cp.add_run(str(v));r.font.name='Times New Roman';r.font.size=Pt(10)

H("CHAPTER 3",16,A.CENTER);H("METHODOLOGY",14,A.CENTER);NL()

H("3.1  Overview of the Proposed Framework")
P("This chapter presents the detailed methodology adopted for the parametric analysis and optimization of bifacial photovoltaic systems. The proposed framework integrates three core components: (i) a mathematical formulation for computing the front-side and rear-side irradiance on bifacial PV modules, (ii) a MATLAB/Simulink model for simulating the electrical characteristics of the PV system, and (iii) a user-oriented frontend platform that retrieves location-specific irradiance data and performs automated parametric sweep analyses.")
P("The framework is designed to evaluate and determine the optimal configuration of bifacial PV systems by varying key installation parameters including module tilt angle, ground surface albedo, and mounting height. The computational pipeline begins with the user specifying the geographical location, date range, and parameter sweep ranges through the frontend interface. The backend application processes these inputs, retrieves hourly irradiance data from the NASA POWER API, and executes the analysis engine to compute the system performance metrics for each combination of installation parameters. The results are presented through an interactive dashboard displaying optimal configurations, I-V and P-V characteristics, and comparative performance metrics.")
P("The mathematical formulation implemented in this framework draws upon established solar radiation models from the literature, specifically the Liu-Jordan isotropic sky model [15] for front-side irradiance estimation and the view factor approach proposed by Yusufoglu et al. [18] for accurate rear-side irradiance calculation with shadow correction. The integration of these models with real-time irradiance data from the NASA POWER database enables location-specific optimization that accounts for the actual solar resource characteristics of the selected site.")

H("3.2  Mathematical Formulation of Bifacial PV Operation")
P("Unlike conventional monofacial PV modules, which receive irradiance only on the front surface, bifacial modules are capable of utilizing additional radiation reflected from the ground and surrounding surfaces. The total effective irradiance on a bifacial module can therefore be expressed as the sum of the front and rear irradiance components [4]:")
EQ("G\u209C\u2092\u209C = G\u2083\u1D63\u2092\u2099\u209C + G\u1D63\u2091\u2090\u1D63", 1)
P("where G_tot is the total irradiance received by the bifacial module (W/m\u00b2), G_front is the irradiance incident on the front surface (W/m\u00b2), and G_rear is the irradiance incident on the rear surface (W/m\u00b2). The front irradiance and rear irradiance depend on several factors which are taken into account in the analysis as described in the following subsections.")

H("3.2.1  Front-Side Irradiance Model",12)
P("The front-side irradiance is estimated using the Liu-Jordan isotropic sky model [15], which decomposes the global horizontal irradiance into beam, diffuse, and ground-reflected components. This model assumes that diffuse radiation from the sky is uniformly distributed across the sky hemisphere, which provides a practical and widely validated approximation for engineering applications.")
P("The direct beam irradiance incident on the tilted module surface is calculated as:")
EQ("beam_tilted = DNI \u00d7 max(0, cos \u03b8\u1d62)", 2)
P("where DNI represents the direct normal irradiance (W/m\u00b2) and \u03b8_i denotes the angle of incidence between the incoming solar radiation and the normal to the module surface. The max(0, cos \u03b8_i) term ensures that negative irradiance values are avoided when the sun is below the panel horizon, which would correspond to a physically impossible scenario of illumination from behind the module.")
P("The diffuse irradiance component is modeled assuming isotropic sky distribution. The sky view factor and ground view factor for the tilted surface are given by [15]:")
EQ("sky_vf = (1 + cos \u03b2) / 2", 3)
EQ("gnd_vf = (1 \u2212 cos \u03b2) / 2", 4)
P("where \u03b2 represents the module tilt angle measured from the horizontal plane. The sky view factor quantifies the portion of the sky dome visible to the tilted PV surface and determines the fraction of diffuse sky radiation intercepted by the module. Similarly, the ground view factor represents the fraction of ground-reflected radiation reaching the module surface. These two factors are complementary, summing to unity for any tilt angle.")
P("The diffuse component of global irradiance is given by DHI (diffuse horizontal irradiance). Combining these contributions, the total front-side irradiance is calculated as [15]:")
EQ("G_front = beam_tilted + (DHI \u00d7 sky_vf) + (GHI \u00d7 \u03c1 \u00d7 gnd_vf)", 5)
P("where GHI is the global horizontal irradiance (W/m\u00b2), DHI is the diffuse horizontal irradiance (W/m\u00b2), and \u03c1 is the ground surface albedo coefficient (dimensionless). The three terms in this equation represent the direct beam, diffuse sky, and ground-reflected components of front-side irradiance, respectively.")

H("3.2.2  Rear-Side Irradiance Model",12)
P("The computation of rear irradiance is more complex than the front-side calculation because it must account for the shadow cast by the module on the ground beneath it. The rear irradiance depends on the shadow view factor F_V, albedo \u03c1, DHI, GHI, and rear sky view factor. The shadow view factor F_V represents the reduction in rear irradiance due to row-to-row shading in the PV array. It is determined by evaluating the solar profile angle and the geometric relationship between the sun position and the module orientation.")
P("The solar profile angle is computed as [16], [17]:")
EQ("\u03b3_profile = tan\u207b\u00b9(tan(\u03b1_sun) / cos(\u03b3_sun \u2212 \u03b3_panel))", 6)
P("where \u03b1_sun is the solar elevation angle, and \u03b3_sun and \u03b3_panel represent the solar and panel azimuth angles, respectively. The solar profile angle provides a convenient way to characterize the apparent position of the sun relative to the plane perpendicular to the module surface.")

H("3.2.3  Shadow and View Factor Computation",12)
P("The lower and upper bounds of the shadow region cast by the module on the ground are calculated using the solar profile angle as follows [16], [17]:")
EQ("shadow_lower = h / tan(\u03b3_profile)", 7)
EQ("shadow_upper = (h + W\u00b7sin \u03b2) / tan(\u03b3_profile)", 8)
P("where h is the mounting height of the lower edge of the module above the ground (m), W is the width of the module (m), and \u03b2 is the tilt angle. These equations define the extent of the shadow region on the ground, which determines the area from which direct beam radiation is blocked and therefore cannot contribute to ground-reflected irradiance reaching the rear surface.")
P("The rear sky view factor (rear_vf) is geometrically represented by the same form as equation (3) with \u03b2 replaced by (180\u00b0 \u2212 \u03b2), which accounts for the reversed orientation of the rear surface relative to the ground. The shadow view factor F_V is calculated by integrating over the shadow region as described by Yusufoglu et al. [18]:")
EQ("F_V = rear_vf \u2212 \u222b cos(\u03b8\u2081)\u00b7cos(\u03b8\u2082) / (2\u00b7r_shadow) dx", 9)
P("where \u03b8\u2081 and \u03b8\u2082 represent the geometric angles between the ground and module surfaces, and r is the distance between the differential ground and module surface elements. This integral is evaluated numerically over the shadow region to obtain the net view factor reduction due to shading.")
FIG("Figure 3.2","View Factor representation for Bifacial PV [18]")
P("The final rear irradiance is given by incorporating the above factors. Following the methodology of Yusufoglu et al. [18], the ground-reflected radiation is separated into diffuse and direct components:")
EQ("G_r0 = \u03c1 \u00b7 DHI \u00b7 rear_vf + \u03c1 \u00b7 (GHI \u2212 DHI) \u00b7 (rear_vf \u2212 F_V)", 10)
P("This formulation recognizes that the diffuse component (DHI) is reflected from the entire ground surface including the shadowed area, since diffuse radiation arrives from all directions and is not blocked by the module shadow. The direct component (GHI \u2212 DHI), however, is reflected only from the unshadowed portion of the ground, hence the subtraction of the shadow view factor F_V from the rear view factor.")
P("The effective rear irradiance is then scaled by the bifaciality factor to account for the difference in conversion efficiency between the front and rear surfaces of the module:")
EQ("G_rear = G_r0 \u00d7 \u03c6_bifaciality", 11)
P("The bifaciality factor \u03c6 represents the ratio of rear-side efficiency to front-side efficiency of the bifacial PV module and is typically in the range of 0.70 to 0.90 for modern bifacial modules. In this study, a bifaciality factor of 0.7 is assumed based on practical considerations and manufacturer specifications for commercially available bifacial modules.")

H("3.3  Simulink Model Development")
P("The bifacial photovoltaic system is modeled in MATLAB/Simulink to analyze the electrical performance of the module under varying environmental and installation conditions. The model incorporates key parameters such as global horizontal irradiance (GHI), ground reflectivity (albedo), tilt angle, mounting height, and ambient temperature. The Simulink model provides a validated platform for computing the I-V and P-V characteristics of the bifacial PV module under any combination of input conditions.")
P("The input parameters are processed through a MATLAB function block (calculate_irradiance), which implements the mathematical formulation described in Section 3.2 to determine the effective irradiance G_eff received by the module, corresponding to G_tot in equation (1). The resulting effective irradiance value is then supplied to the PV Array block as the irradiance input. The PV Array block represents the electrical model of the photovoltaic module based on the single-diode equivalent circuit and produces the output current and voltage corresponding to the applied operating conditions.")
P("To obtain the electrical characteristics of the PV system, a controlled voltage source is used to perform an I-V sweep across the PV module terminals. The voltage sweep is generated using a ramp signal (V_ramp), which gradually varies the applied voltage across the PV array over the simulation interval from 0 V to the open-circuit voltage. As the voltage changes, the PV module produces corresponding current values, allowing the complete I-V characteristic curve to be obtained. The measured current (I_meas) and voltage (V_meas) signals are collected through a bus creator block and forwarded to subsequent processing stages. The first scope plots the current-voltage (I-V) characteristics, while the second scope plots the power-voltage (P-V) characteristics.")
FIG("Figure 3.3","Simulink model for Bifacial PV Module")

TBL(["Property","Value"],
    [["Module Type","Bifacial Monocrystalline Silicon"],["Nominal Power","550 Wp (front side)"],["Bifaciality Factor","0.7"],["Open Circuit Voltage (Voc)","49.90 V"],["Short Circuit Current (Isc)","14.0 A"],["Number of Cells","72 x 2 (half-cell)"],["Module Width","1134 mm"],["Module Length","2278 mm"],["NOCT","45 \u00b0C"],["Temperature Coefficient (Voc)","-0.27 %/\u00b0C"]],
    "Table 3.1: Technical Specifications of the PV Module Used in Simulation")
NL()

H("3.4  Frontend Platform Architecture")
P("The proposed framework is designed as a multi-layered system consisting of four primary components: the user interface, backend application, analysis and simulation engine, and the results dashboard. This architecture ensures modularity, maintainability, and extensibility of the platform.")

H("3.4.1  User Interface Layer",12)
P("The process begins at the user interface, where the user provides the necessary input parameters required for the analysis. These inputs include the geographical location in the form of a selected city, the desired date range for analysis, and key system configuration parameters such as module tilt angle range, mounting height range, and ground surface type. The interface is designed to be intuitive and accessible, requiring no specialized knowledge of PV system modeling or solar engineering. Users can select from predefined surface types with known albedo values or specify custom albedo coefficients. The interface also provides the option to select the optimization objective: maximum total energy or maximum rear-side gain.")

H("3.4.2  Backend Application Layer",12)
P("Once the input parameters are specified, the backend application processes these inputs and generates a request for solar irradiance data. The irradiance data is obtained through the NASA POWER API, which provides location-specific solar radiation data for the selected city and time period. The backend handles API communication, data parsing, error handling, and the coordination of parametric sweep computations. The backend is implemented using Python with Flask for the web server component, ensuring cross-platform compatibility and ease of deployment.")

H("3.4.3  Analysis and Simulation Engine",12)
P("The irradiance dataset is passed to the analysis and simulation engine, which performs the computational evaluation of the bifacial PV system. This module integrates the irradiance data with the user-defined installation parameters to simulate the electrical behavior of the bifacial PV system using the mathematical formulation described in Section 3.2. The simulation process evaluates the I-V and P-V characteristics and calculates key performance indicators such as output power, total energy, and bifacial gain under different configuration settings. The engine performs the parametric sweep across all specified combinations of tilt angle, albedo, and mounting height, computing performance metrics for each evaluation point.")

H("3.4.4  Results Dashboard",12)
P("The computed results are presented through a results dashboard providing structured visualization of system performance. The dashboard displays maximum power output and bifacial gain values in tabular form, along with graphical representations of the I-V and P-V characteristics. The dashboard also identifies and highlights the optimal configuration based on the user selected optimization objective (maximum energy or maximum rear gain). Additionally, ranking tables showing the top configurations are provided to enable comparative assessment of different installation options.")
FIG("Figure 3.4","Proposed Framework for Optimal Configuration of Bifacial PV Systems")

H("3.5  NASA POWER API Integration")
P("The NASA Prediction of Worldwide Energy Resources (POWER) project provides access to meteorological and solar radiation data derived from satellite observations and atmospheric models. The POWER API serves as the primary data source for location-specific irradiance information in the proposed framework. For a given geographic location specified by latitude and longitude coordinates, the API returns hourly values of GHI, DHI, DNI, temperature, and other relevant meteorological parameters. The data is available for any location worldwide with a spatial resolution of 0.5 degrees, making it suitable for site-specific analysis across diverse geographic regions.")
P("The API is accessed through HTTP GET requests with parameters specifying the location coordinates, date range, and desired data variables. The returned JSON data is parsed and processed by the backend application to extract the required irradiance components for the bifacial PV performance analysis. Error handling mechanisms are implemented to manage network connectivity issues, API rate limits, and data quality checks.")

H("3.6  Parametric Sweep Configuration")
P("The parametric sweep is configured to systematically evaluate a grid of installation parameter combinations. In the present implementation, the following parameter ranges and increments are used: module mounting height is varied from 50 cm to 450 cm with increments of 50 cm (9 levels); tilt angle is varied from 10 degrees to 50 degrees in steps of 10 degrees (5 levels); and surface albedo is selected from predefined surface types with known albedo values. This configuration results in 45 evaluation points (9 heights x 5 tilt angles) for each selected albedo value.")
P("For each evaluation point, the analysis engine computes the total energy output, peak power, and rear gain percentage. The results are ranked according to the selected optimization objective, and the top configurations are presented to the user. Additionally, the framework provides a provision for fixing two parameters and varying the third to obtain a focused analysis of individual parameter effects. This makes the platform dynamic and adaptable to the practical needs of PV system designers who may wish to explore the sensitivity of performance to specific parameters while keeping others constant.")
FIG("Figure 3.5","Frontend user interface showing parameter sweep configuration")

doc.add_page_break()
doc.save('_part3.docx')
print("Ch3 done.")

"""
Part 1: Generate the BTech Project Report - Front matter and Chapters 1-2
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# PAGE SETUP
# ============================================================
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# Helper functions
def add_heading_custom(doc, text, level=1, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, spacing_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    return p

def add_empty_lines(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

# ============================================================
# TITLE PAGE
# ============================================================
add_empty_lines(doc, 3)
add_heading_custom(doc, 'B.Tech. Project', level=1, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading_custom(doc, 'On', level=1, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 1)
add_heading_custom(doc, 'OPTIMAL CONFIGURATION OF BIFACIAL\nPHOTOVOLTAIC MODULES USING\nPARAMETRIC ANALYSIS', level=1, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 1)
add_para(doc, 'Report submitted in partial fulfillment of the requirements for the\nB. Tech. degree in Electrical Engineering', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_empty_lines(doc, 1)
add_heading_custom(doc, 'By', level=1, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

# Student table
table = doc.add_table(rows=5, cols=2)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
students = [
    ('Name of Student', 'Roll No.'),
    ('Ashish Upadhyay', '2022UEE4521'),
    ('Pratyai Chakrabarty', '2022UEE4586'),
    ('Aman', '2022UEE4532'),
    ('Aditya Garg', '2022UEE4503'),
]
for i, (name, roll) in enumerate(students):
    for j, text in enumerate([name, roll]):
        cell = table.cell(i, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if i == 0:
            run.bold = True

add_empty_lines(doc, 1)
add_para(doc, 'Under the supervision of', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_heading_custom(doc, 'Dr. Astitva Kumar', level=1, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 2)

# Add NSUT logo placeholder text
add_para(doc, '[NSUT Logo]', align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
add_empty_lines(doc, 1)
add_heading_custom(doc, 'DEPARTMENT OF ELECTRICAL ENGINEERING', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading_custom(doc, 'NETAJI SUBHAS UNIVERSITY OF TECHNOLOGY (NSUT)\nDWARKA, NEW DELHI', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading_custom(doc, 'MAY 2026', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# CANDIDATE(S) DECLARATION (page i)
# ============================================================
add_heading_custom(doc, "CANDIDATE(S) DECLARATION", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 1)
add_heading_custom(doc, "DEPARTMENT OF ELECTRICAL ENGINEERING", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 1)

decl_text = (
    'I/We, Ashish Upadhyay (2022UEE4521), Pratyai Chakrabarty (2022UEE4586), '
    'Aman (2022UEE4532) and Aditya Garg (2022UEE4503), students of B. Tech., '
    'Department of Electrical Engineering, hereby declare that the Project-Thesis titled '
    '"Optimal Configuration of Bifacial Photovoltaic Modules using Parametric Analysis" '
    'which is submitted by us to the Department of Electrical Engineering, Netaji Subhas '
    'University of Technology (NSUT) Dwarka, New Delhi in partial fulfillment of the '
    'requirement for the award of the degree of Bachelor of Technology is our original work '
    'and not copied from any source without proper citation. The manuscript has been '
    'subjected to plagiarism check by Turnitin software. This work has not previously '
    'formed the basis for the award of any other Degree.'
)
add_para(doc, decl_text)
add_empty_lines(doc, 2)
add_para(doc, 'Place: New Delhi')
add_para(doc, 'Date: May 2026')
add_empty_lines(doc, 2)

for name, roll in students[1:]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f'{name}\n{roll}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

add_para(doc, '\ni', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_page_break()

# ============================================================
# CERTIFICATE OF DECLARATION (page ii)
# ============================================================
add_heading_custom(doc, "CERTIFICATE OF DECLARATION", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 1)
add_heading_custom(doc, "DEPARTMENT OF ELECTRICAL ENGINEERING", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 1)

cert_text = (
    'This is to certify that the work embodied in project thesis titled, '
    '"Optimal Configuration of Bifacial Photovoltaic Modules using Parametric Analysis" '
    'by Ashish Upadhyay (2022UEE4521), Pratyai Chakrabarty (2022UEE4586), '
    'Aman (2022UEE4532) and Aditya Garg (2022UEE4503) is the bonafide work of the group '
    'submitted to Netaji Subhas University of Technology for consideration in 8th Semester '
    'B.Tech. Project Evaluation.'
)
add_para(doc, cert_text)
add_empty_lines(doc, 1)

cert_text2 = (
    'The original Research work was carried out by the team under my guidance and supervision '
    'in the academic year 2025-2026. This work has not been submitted for any other diploma or degree '
    'of any university. On the basis of declaration made by the group, we recommend the project report '
    'for evaluation.'
)
add_para(doc, cert_text2)
add_empty_lines(doc, 3)

add_para(doc, 'Dr. Astitva Kumar')
add_para(doc, 'Assistant Professor')
add_para(doc, 'Department of Electrical Engineering')
add_para(doc, 'Netaji Subhas University of Technology')
add_para(doc, '\nii', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_page_break()

# ============================================================
# ACKNOWLEDGEMENT (page iii)
# ============================================================
add_heading_custom(doc, "ACKNOWLEDGEMENT", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 1)

ack1 = (
    'We would like to express our sincere gratitude and appreciation to all those who made it '
    'possible to complete this project. First and foremost, we extend our deepest thanks to our '
    'project supervisor, Dr. Astitva Kumar, Assistant Professor, Department of Electrical Engineering, '
    'Netaji Subhas University of Technology, New Delhi, whose continuous guidance, invaluable suggestions, '
    'constructive criticism, and constant encouragement have been instrumental throughout the course of '
    'this work. His expertise in the domain of photovoltaic systems and renewable energy provided the '
    'foundation upon which this research was built.'
)
add_para(doc, ack1)

ack2 = (
    'We would also like to acknowledge the Department of Electrical Engineering, NSUT, for providing '
    'access to laboratory facilities, computational resources, and licensed software tools including '
    'MATLAB/Simulink, which were essential for the simulation and analysis components of this project. '
    'The department\'s commitment to fostering a research-oriented environment has been greatly appreciated.'
)
add_para(doc, ack2)

ack3 = (
    'We are grateful to our colleagues and peers who offered their time for technical discussions, '
    'proofreading, and constructive feedback during the preparation of this report. Their support '
    'has been invaluable in refining both the technical content and the presentation quality of this work.'
)
add_para(doc, ack3)

ack4 = (
    'Finally, we wish to express our heartfelt gratitude to our families for their unwavering support, '
    'patience, and motivation throughout the duration of this project. Their encouragement has been a '
    'constant source of strength.'
)
add_para(doc, ack4)

add_empty_lines(doc, 2)
for name, roll in students[1:]:
    add_para(doc, name, align=WD_ALIGN_PARAGRAPH.RIGHT)

add_para(doc, '\niii', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_page_break()

# ============================================================
# PLAGIARISM REPORT (page iv)
# ============================================================
add_heading_custom(doc, "PLAGIARISM REPORT", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 2)
add_para(doc, '[Attach Turnitin/Ouriginal plagiarism report here]', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_para(doc, 'Similarity Index: ___ % (must be less than 20%)', align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 2)
add_para(doc, 'Verified by:', align=WD_ALIGN_PARAGRAPH.LEFT)
add_empty_lines(doc, 2)
add_para(doc, 'Dr. Astitva Kumar')
add_para(doc, 'Supervisor')
add_para(doc, '\niv', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_page_break()

# ============================================================
# ABSTRACT (page v)
# ============================================================
add_heading_custom(doc, "ABSTRACT", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 1)

abstract = (
    'Solar photovoltaic (PV) technology has become a key component in the global transition toward '
    'sustainable energy, with countries such as India witnessing rapid growth in solar power deployment '
    'due to abundant solar resources and increasing energy demand. Efficient utilization of available '
    'solar irradiance is therefore critical to maximize energy generation from installed PV systems. '
    'In this context, bifacial photovoltaic modules have gained significant attention as they are capable '
    'of capturing solar irradiance from both the front and rear surfaces, leading to higher energy yield '
    'compared to conventional monofacial panels. However, the performance advantage of bifacial PV systems '
    'strongly depends on installation parameters such as module tilt angle, ground surface albedo, and '
    'mounting height, which influence the amount of reflected and incident irradiance received by the rear '
    'side of the module. This project presents a computational framework for determining the optimal '
    'configuration of bifacial PV installations for a given geographical location. A user-oriented front-end '
    'platform is developed where inputs such as city and date are provided, and corresponding hourly irradiance '
    'data are retrieved from the NASA POWER database. Parametric analyses are then performed by varying key '
    'installation parameters including tilt angle, surface albedo, and mounting height. The framework evaluates '
    'the resulting energy output and identifies the configuration that maximizes power generation for the selected '
    'location. The proposed approach enables location-specific optimization of bifacial PV systems and provides '
    'a practical tool for improving system design and deployment strategies.'
)
add_para(doc, abstract)
add_empty_lines(doc, 1)
add_para(doc, 'Keywords: Solar Photovoltaic, Bifacial PV, Parametric Analysis, Frontend Platform, Strategic Installation, View Factor', bold=True)
add_para(doc, '\nv', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (pages vi-viii)
# ============================================================
add_heading_custom(doc, "TABLE OF CONTENTS", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 1)

toc_entries = [
    ("CANDIDATE(S) DECLARATION", "i", 0),
    ("CERTIFICATE OF DECLARATION", "ii", 0),
    ("ACKNOWLEDGEMENT", "iii", 0),
    ("PLAGIARISM REPORT", "iv", 0),
    ("ABSTRACT", "v", 0),
    ("TABLE OF CONTENTS", "vi", 0),
    ("LIST OF FIGURES", "ix", 0),
    ("LIST OF TABLES", "x", 0),
    ("LIST OF ABBREVIATIONS", "xi", 0),
    ("", "", 0),
    ("CHAPTER 1: INTRODUCTION", "1", 0),
    ("1.1  Background and Motivation", "1", 1),
    ("1.2  Solar Photovoltaic Technology: An Overview", "3", 1),
    ("1.3  Monofacial versus Bifacial PV Modules", "4", 1),
    ("1.4  Key Parameters Affecting Bifacial PV Performance", "5", 1),
    ("1.5  Challenges in Bifacial PV Optimization", "7", 1),
    ("1.6  Problem Statement and Research Objectives", "8", 1),
    ("1.7  Organization of the Report", "9", 1),
    ("", "", 0),
    ("CHAPTER 2: LITERATURE REVIEW", "10", 0),
    ("2.1  Review of Bifacial PV Technology", "10", 1),
    ("2.2  Influence of Tilt Angle on Bifacial PV Performance", "12", 1),
    ("2.3  Effect of Ground Albedo on Rear-Side Irradiance", "13", 1),
    ("2.4  Impact of Mounting Height on Energy Yield", "15", 1),
    ("2.5  Combined Parametric Optimization Studies", "16", 1),
    ("2.6  Simulation and Modeling Approaches", "17", 1),
    ("2.7  Research Gap Analysis", "18", 1),
    ("", "", 0),
    ("CHAPTER 3: METHODOLOGY", "20", 0),
    ("3.1  Overview of the Proposed Framework", "20", 1),
    ("3.2  Mathematical Formulation of Bifacial PV Operation", "21", 1),
    ("    3.2.1  Front-Side Irradiance Model", "22", 2),
    ("    3.2.2  Rear-Side Irradiance Model", "23", 2),
    ("    3.2.3  Shadow and View Factor Computation", "24", 2),
    ("3.3  Simulink Model Development", "26", 1),
    ("3.4  Frontend Platform Architecture", "28", 1),
    ("    3.4.1  User Interface Layer", "28", 2),
    ("    3.4.2  Backend Application Layer", "29", 2),
    ("    3.4.3  Analysis and Simulation Engine", "29", 2),
    ("    3.4.4  Results Dashboard", "30", 2),
    ("3.5  NASA POWER API Integration", "30", 1),
    ("3.6  Parametric Sweep Configuration", "31", 1),
    ("", "", 0),
    ("CHAPTER 4: RESULTS AND DISCUSSION", "32", 0),
    ("4.1  Case Study Configuration", "32", 1),
    ("4.2  Combined Parametric Sweep Results", "33", 1),
    ("4.3  I-V and P-V Characteristics", "34", 1),
    ("4.4  Optimal Configuration: Maximum Energy", "35", 1),
    ("4.5  Optimal Configuration: Maximum Rear Gain", "36", 1),
    ("4.6  Individual Parameter Variation Analysis", "37", 1),
    ("    4.6.1  Effect of Albedo Variation", "37", 2),
    ("    4.6.2  Effect of Height Variation", "38", 2),
    ("    4.6.3  Effect of Tilt Angle Variation", "39", 2),
    ("4.7  Comparison with Published Literature", "40", 1),
    ("4.8  Discussion of Results", "41", 1),
    ("", "", 0),
    ("CHAPTER 5: CONCLUSION AND FUTURE WORK", "43", 0),
    ("5.1  Summary of Contributions", "43", 1),
    ("5.2  Key Findings", "44", 1),
    ("5.3  Limitations", "45", 1),
    ("5.4  Scope for Future Work", "45", 1),
    ("", "", 0),
    ("REFERENCES", "47", 0),
    ("APPENDIX A: Software Code", "50", 0),
    ("APPENDIX B: Simulink Block Diagrams", "52", 0),
    ("BIO-DATA", "54", 0),
]

for entry, page, level in toc_entries:
    if entry == "":
        add_empty_lines(doc, 0)
        continue
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    indent = "    " * level
    run = p.add_run(f'{indent}{entry}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if level == 0:
        run.bold = True
    # Add tab and page number
    tab_run = p.add_run(f'\t{page}')
    tab_run.font.name = 'Times New Roman'
    tab_run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# LIST OF FIGURES
# ============================================================
add_heading_custom(doc, "LIST OF FIGURES", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 1)

figures = [
    ("Figure 1.1", "Global solar PV installed capacity growth (2015-2025)", "2"),
    ("Figure 1.2", "Schematic representation of monofacial and bifacial PV modules", "4"),
    ("Figure 1.3", "Key parameters influencing bifacial PV module performance", "6"),
    ("Figure 2.1", "Year-wise publication trend on bifacial PV research", "11"),
    ("Figure 2.2", "Albedo values of common ground surfaces for PV applications", "14"),
    ("Figure 3.1", "Key parameters affecting bifacial PV modules", "21"),
    ("Figure 3.2", "View factor representation for bifacial PV", "25"),
    ("Figure 3.3", "Simulink model for bifacial PV module", "27"),
    ("Figure 3.4", "Proposed framework for optimal configuration of bifacial PV systems", "29"),
    ("Figure 4.1", "I-V characteristics for different configurations", "34"),
    ("Figure 4.2", "P-V characteristics for different configurations", "35"),
    ("Figure 4.3", "Frontend interface for parameter sweep analysis", "37"),
    ("Figure 4.4", "Rear share of effective irradiance on varying albedo", "38"),
    ("Figure 4.5", "Rear share of effective irradiance on varying height", "39"),
    ("Figure 4.6", "Rear share of effective irradiance on varying tilt angle", "40"),
]

for fig_num, caption, page in figures:
    p = doc.add_paragraph()
    run = p.add_run(f'{fig_num}: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(f'{caption}')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run3 = p.add_run(f'\t{page}')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)

add_para(doc, '\nix', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_page_break()

# ============================================================
# LIST OF TABLES
# ============================================================
add_heading_custom(doc, "LIST OF TABLES", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 1)

tables_list = [
    ("Table 2.1", "Albedo values of common surface types for PV applications", "14"),
    ("Table 2.2", "Summary of key literature on bifacial PV parametric studies", "18"),
    ("Table 3.1", "Technical specifications of the PV module used in simulation", "26"),
    ("Table 4.1", "Max energy based optimal configuration (Top 5)", "36"),
    ("Table 4.2", "Max rear gain based optimal configuration (Top 5)", "36"),
    ("Table 4.3", "Comparison of results with published literature", "41"),
]

for tab_num, caption, page in tables_list:
    p = doc.add_paragraph()
    run = p.add_run(f'{tab_num}: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(f'{caption}')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run3 = p.add_run(f'\t{page}')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)

add_para(doc, '\nx', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_page_break()

# ============================================================
# LIST OF ABBREVIATIONS
# ============================================================
add_heading_custom(doc, "LIST OF ABBREVIATIONS", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(doc, 1)

abbreviations = [
    ("PV", "Photovoltaic"),
    ("BPV", "Bifacial Photovoltaic"),
    ("GHI", "Global Horizontal Irradiance"),
    ("DHI", "Diffuse Horizontal Irradiance"),
    ("DNI", "Direct Normal Irradiance"),
    ("API", "Application Programming Interface"),
    ("NASA", "National Aeronautics and Space Administration"),
    ("POWER", "Prediction of Worldwide Energy Resources"),
    ("NOCT", "Nominal Operating Cell Temperature"),
    ("STC", "Standard Test Conditions"),
    ("I-V", "Current-Voltage"),
    ("P-V", "Power-Voltage"),
    ("MATLAB", "Matrix Laboratory"),
    ("IEEE", "Institute of Electrical and Electronics Engineers"),
    ("NSUT", "Netaji Subhas University of Technology"),
    ("MPP", "Maximum Power Point"),
]

for abbr, full in abbreviations:
    p = doc.add_paragraph()
    run = p.add_run(f'{abbr}')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(f'\t\t{full}')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

add_para(doc, '\nxi', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_page_break()

# Save intermediate
doc.save('BTP_Report_Part1.docx')
print("Part 1 saved successfully!")
